
# docx_processor.py

import re
import base64
from io import BytesIO
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.ns import qn
from docx.table import Table as DocxTable, _Cell
from docx.table import Table 
from docx.text.paragraph import Paragraph
from xml.etree.ElementTree import Element, SubElement, tostring
from xml.dom import minidom
# from tinhoc_processor import TinHocProcessor # Bỏ import nếu chưa có
from typing import List, Union, Any, Optional
import traceback
from PIL import Image
from io import BytesIO
from bs4 import BeautifulSoup


try:
    from tinhoc_processor import TinHocProcessor
except ImportError:
    class TinHocProcessor:
        def __init__(self): pass
        def dang_ds_tinhoc(self, cau_sau_xu_ly, xml, audio, doc): pass
        def dang_tn_tinhoc(self, cau_sau_xu_ly, xml, audio, doc): pass
        def dang_dt(self, cau_sau_xu_ly, xml, subject): pass
        def dang_tl(self, cau_sau_xu_ly, xml, audio): pass


class DocxProcessor:
    """Class chính xử lý DOCX"""
    def __init__(self):
        self.subjects_with_default_titles = [
            "TOANTHPT", "VATLITHPT2", "HOATHPT2", "SINHTHPT2",
            "LICHSUTHPT", "DIALITHPT", "GDCDTHPT2", "NGUVANTHPT","VATLYTHPT2",
            "TOANTHCS2", "KHTN", "KHXHTHCS", "GDCDTHCS2", "NGUVANTHCS2", "DGNLDHQGHN","DETHI","CAMBRIDGE"
        ]
        self.tinhoc_subjects = ['TINHOCTHPT', 'TINHOC3']
        self.index_question = 0
        self.tinhoc_processor = TinHocProcessor()
        self.nsmap = {
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'v': 'urn:schemas-microsoft-com:vml',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    }

   

    def process_docx(self, file_path):
        """Xử lý file DOCX và trả về XML string hoặc danh sách lỗi"""
        errors = []
        doc = None
        
        try:
            print(f">>>>> Debug file path {file_path}")
            doc = Document(file_path)
            self.doc = doc
            self.tinhoc_processor.doc = self.doc
            body = doc.element.body
            
            # Parse các elements theo thứ tự trong body
            paragraphs = []
            try:
                for child in body:
                    if isinstance(child, CT_P):
                        paragraphs.append(Paragraph(child, doc))
                    elif isinstance(child, CT_Tbl):
                        paragraphs.append(Table(child, doc))
            except Exception as e:
                errors.append(f"Lỗi khi đọc cấu trúc body của DOCX: {str(e)}")
                return "", errors
            
            # Biến trạng thái
            list_hl = []
            group_of_questions = []
            current_tag = None
            current_table = None
            content_hl = False
            
            for idx, para in enumerate(paragraphs):
                try:
                    is_table = isinstance(para, Table)
                    
                    # Xử lý table
                    if is_table:
                        current_table = para
                        
                        # ✅ SỬA: Thêm table vào học liệu nếu đang trong chế độ HL
                        if content_hl and list_hl:
                            list_hl[-1]['content'].append(current_table)
                            print(f"[DEBUG] ✓ Thêm table vào học liệu tại idx={idx}")
                            continue
                        
                        # Thêm vào câu hỏi thường
                        if group_of_questions and group_of_questions[-1]['questions']:
                            group_of_questions[-1]['questions'].append(current_table)
                        continue
                    
                    # Bỏ qua paragraph rỗng
                    if len(para.runs) == 0:
                        continue
                    
                    text = para.text.strip()
                    
                    # ——— ƯU TIÊN 1: XỬ LÝ HEADER [tag, posttype, level] ———
                    if re.match(r'^\[.*\]$', text):
                        header = text.replace('[', '').replace(']', '')
                        fields = [f.strip() for f in header.split(',')]
                        
                        if len(fields) != 3:
                            errors.append(f"Sai format header tại dòng {idx + 1}: {text}")
                            continue
                        
                        dvkt, posttype, knowledge = fields
                        current_tag = dvkt
                        cap_do = ['NB', 'TH', 'VD', 'VDC']
                        knowledge_upper = knowledge.upper()
                        level = cap_do.index(knowledge_upper) if knowledge_upper in cap_do else 0
                        
                        group = {
                            'subject': dvkt.split('_')[0],
                            'tag': dvkt,
                            'original_tag': dvkt,
                            'posttype': posttype,
                            'knowledgelevel': knowledge_upper if knowledge_upper in cap_do else 'NB',
                            'level': level,
                            'questions': []
                        }
                        group_of_questions.append(group)
                        content_hl = False
                        continue
                    
                    # ——— ƯU TIÊN 2: XỬ LÝ DÒNG BẮT ĐẦU BẰNG "HL:" ———
                    if text.startswith('HL:'):
                        if list_hl:
                            prev_group = group_of_questions[-1]
                            group_of_questions = [{
                                'subject': prev_group['subject'],
                                'tag': prev_group['tag'],
                                'posttype': prev_group['posttype'],
                                'knowledgelevel': prev_group['knowledgelevel'],
                                'level': prev_group['level'],
                                'questions': []
                            }]
                        
                        hoc_lieu = {
                            'content': [para],  # Bắt đầu với paragraph "HL:"
                            'groupOfQ': group_of_questions
                        }
                        content_hl = True
                        list_hl.append(hoc_lieu)
                        print(f"[DEBUG] ✓ Tạo học liệu mới tại idx={idx}")
                        continue
                    
                    # ——— ƯU TIÊN 3: PHÁT HIỆN CÂU HỎI MỚI ———
                    if re.match(r'^C[âa]u\s*\d', text, re.IGNORECASE):
                        content_hl = False
                    
                    # ——— THÊM VÀO NỘI DUNG HỌC LIỆU (NẾU ĐANG TRONG CHẾ ĐỘ HL) ———
                    if content_hl and list_hl:
                        list_hl[-1]['content'].append(para)
                        print(f"[DEBUG] ✓ Thêm paragraph vào học liệu tại idx={idx}")
                        continue
                    
                    # ——— THÊM VÀO CÂU HỎI THƯỜNG ———
                    if group_of_questions:
                        para.current_tag = current_tag
                        group_of_questions[-1]['questions'].append(para)
                        
                except Exception as e:
                    import traceback
                    errors.append(f"Lỗi khi xử lý paragraph #{idx} (text: {getattr(para, 'text', 'N/A')[:50]}...): {str(e)}")
                    continue
            
            # Tạo XML
            try:
                if list_hl:
                    root = Element('itemDocuments')
                    for idx_hl, hoc_lieu in enumerate(list_hl):
                        print(f"[DEBUG] Xử lý học liệu #{idx_hl}, số phần tử content: {len(hoc_lieu['content'])}")
                        item_doc = self.create_hoc_lieu_xml(hoc_lieu, idx_hl)
                        root.append(item_doc)
                else:
                    root = Element('questions')
                    self.index_question = 0
                    for group in group_of_questions:
                        self.format_questions(group, root, errors)
            except Exception as e:
                errors.append(f"Lỗi khi tạo XML: {str(e)}")
                return "", errors
            
            try:
                xml_str = self.prettify_xml(root)
                xml_str = self.post_process_xml(xml_str)
            except Exception as e:
                errors.append(f"Lỗi khi định dạng XML: {str(e)}")
                return "", errors
            
            return xml_str, errors
            
        except Exception as e:
            errors.append(f"Lỗi nghiêm trọng khi xử lý file '{file_path}': {str(e)}")
            import traceback
            traceback.print_exc()
            return "", errors

    def create_hoc_lieu_xml(self, hoc_lieu, index_hl):
        """Tạo XML cho học liệu"""
        item_doc = Element('itemDocument')

        questions_hl = [g for g in hoc_lieu['groupOfQ'] if g['questions']]

        sub_id = SubElement(item_doc, 'subjectId')

        sub_id.text = questions_hl[0]['subject'] if questions_hl else ''

        know_id = SubElement(item_doc, 'knowledgeId')

        know_id.text = questions_hl[0]['tag'] if questions_hl else ''

        group_material = SubElement(item_doc, 'groupQuestionMaterial')

        group_material.text = str(index_hl)

        content_html = SubElement(item_doc, 'contentHtml')

        html_content = self.xu_ly_hl(hoc_lieu['content'])

        content_html.text = html_content

        list_question = SubElement(item_doc, 'listQuestion')
        for group in questions_hl:
            # Gọi format_questions với danh sách lỗi
            self.format_questions(group, list_question, [])
        return item_doc
    

 
    def get_indent_html(self, paragraph: Paragraph):
        """
        Trả về chuỗi thụt lề trái bằng entity HTML.
        Giả sử paragraph.paragraph_format.left_indent trả về giá trị EMU (do lỗi hoặc custom),
        hoặc pt (tiêu chuẩn). Ta phát hiện và xử lý tự động.
        """
        try:
            left_indent = paragraph.paragraph_format.left_indent or 0
            first_line = paragraph.paragraph_format.first_line_indent or 0

            # Chuyển sang số thực
            left_val = float(left_indent) if left_indent else 0.0
            first_val = float(first_line) if first_line else 0.0

            # PHÁT HIỆN: nếu giá trị > 10000 → rất có thể là EMU
            if left_val > 10000:
                # Chuyển EMU → pt
                left_val = left_val / 12700.0
            if first_val > 10000:
                first_val = first_val / 12700.0

            total_pt = left_val + max(0.0, first_val)
        except (AttributeError, TypeError, ValueError):
            total_pt = 0.0

        if total_pt <= 0:
            return ""

        # GAS giả định: 1pt = 1px
        px = int(round(total_pt))

        emsp = px // 16
        px %= 16
        ensp = px // 8
        px %= 8
        thinsp = px // 4
        px %= 4
        hairsp = px // 2

        return "&emsp;" * emsp + "&ensp;" * ensp + "&thinsp;" * thinsp + "&hairsp;" * hairsp
    
    def get_alignment_style(self, paragraph: Paragraph) -> Optional[str]:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        alignment = paragraph.alignment
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return "center"
        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return "right"
        elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return "justify"
        else:
            return None 

 
    def xu_ly_hl(self, content):
        """
        Xử lý nội dung học liệu (HL) thành HTML hoàn chỉnh.
        ✅ ĐÃ SỬA: Phát hiện table đúng cách
        """
        print("[DEBUG] === BẮT ĐẦU HÀM xu_ly_hl ===")
        
        # =================== HELPER: EXTRACT ELEMENTS =================== 
        def extract_elements(container: Any) -> List[Union[Paragraph, DocxTable]]:
            elements = []
            print(f"[DEBUG] extract_elements: container={type(container)}")
            
            try:
                if hasattr(container, "paragraphs") or hasattr(container, "tables"):
                    body_elem = getattr(container, "_element", None)
                    if body_elem is None and hasattr(container, "_body"):
                        body_elem = getattr(container._body, "_element", None)
                    
                    if body_elem is not None:
                        for child in body_elem.iterchildren():
                            # ✅ SỬA: Kiểm tra CT_Tbl thay vì Table
                            if isinstance(child, CT_P):
                                para = Paragraph(child, container)
                                elements.append(para)
                            elif isinstance(child, CT_Tbl):  # ← SỬA ĐÂY
                                tbl = DocxTable(child, container)
                                elements.append(tbl)
                                print(f"[DEBUG] ✓ Phát hiện table trong HL")
                        
                        print(f"[DEBUG] Trích xuất từ XML body: {len(elements)} phần tử")
                        return elements
                    else:
                        paragraphs = list(getattr(container, "paragraphs", []))
                        tables = list(getattr(container, "tables", []))
                        elements = paragraphs + tables
                        print("[WARN] Không xác định được body element, nối thẳng paragraphs+tables")
                        return elements
            except Exception as e:
                print(f"[ERROR] extract_elements lỗi: {e}")
                traceback.print_exc()
                return elements

        # =================== CHUẨN BỊ DANH SÁCH PHẦN TỬ ===================
        if isinstance(content, list):
            all_elements = content
            print(f"[DEBUG] Đầu vào là list, số phần tử: {len(all_elements)}")
        elif hasattr(content, "_element"):
            all_elements = extract_elements(content)
            print(f"[DEBUG] Đầu vào là document/body, trích xuất {len(all_elements)} phần tử")
        else:
            print(f"[WARN] Loại đầu vào không hỗ trợ: {type(content)}")
            return ""
        
        # =================== 🔧 CHUẨN HÓA PHẦN TỬ ===================
        normalized_elements = []
        for el in all_elements:
            if isinstance(el, CT_P):
                normalized_elements.append(Paragraph(el, self.doc))
            elif isinstance(el, CT_Tbl):  # ← SỬA ĐÂY
                normalized_elements.append(DocxTable(el, self.doc))
                print(f"[DEBUG] ✓ Chuẩn hóa table thành DocxTable")
            elif isinstance(el, (Paragraph, DocxTable)):
                normalized_elements.append(el)
            else:
                print(f"[WARN] Bỏ qua phần tử không hỗ trợ trong HL: {type(el)}")
        
        all_elements = normalized_elements
        
        # =================== XÂY DỰNG FRAGMENTS ===================
        fragments = []
        for i, el in enumerate(all_elements):
            print(f"[DEBUG] --- Xử lý phần tử {i}: {type(el).__name__}")
            
            if isinstance(el, DocxTable):
                table_html = self.convert_table_to_html(el, is_hoc_lieu=True)
                fragments.append({
                    'type': 'plain',
                    'alignment': None,
                    'content': table_html
                })
                print(f"[DEBUG] ✓ Đã convert table sang HTML")
            elif isinstance(el, Paragraph):
                align = self.get_alignment_style(el)
                para_html = self.convert_paragraph_for_hl(el)
                
                if para_html.endswith(' '):
                    para_html = para_html[:-5]
                
                if align in ("center", "right", "justify"):
                    fragments.append({
                        'type': 'aligned',
                        'alignment': align,
                        'content': para_html
                    })
                else:
                    fragments.append({
                        'type': 'plain',
                        'alignment': None,
                        'content': para_html
                    })
            else:
                print(f"[WARN] Bỏ qua phần tử loại: {type(el)}")
        
        # =================== GOM NHÓM VÀ RENDER ===================
        result_parts = []
        i = 0
        while i < len(fragments):
            frag = fragments[i]
            if frag['type'] == 'aligned':
                current_align = frag['alignment']
                group_contents = []
                j = i
                while (j < len(fragments) and
                    fragments[j]['type'] == 'aligned' and
                    fragments[j]['alignment'] == current_align):
                    group_contents.append(fragments[j]['content'])
                    j += 1
                # Ghép nội dung, nhưng đảm bảo giữa các phần tử có <br>
                grouped_html = f'<div style="text-align:{current_align}">{" ".join(group_contents)}</div>'
                result_parts.append(grouped_html)
                if j < len(fragments):
                    result_parts.append(' ')
                i = j
            else:
                # Xử lý plain content (có thể là <br/> từ paragraph rỗng)
                content = frag['content']
                result_parts.append(content)
                i += 1

        # ✅ XỬ LÝ NHIỀU <br/> LIÊN TIẾP: chuyển "<br/><br/>" thành đúng 2 dòng
        html = "".join(result_parts)
        print("[DEBUG] === KẾT THÚC HÀM xu_ly_hl ===")
        return html


    def convert_paragraph_for_hl(self, p: Paragraph) -> str:
            """Xử lý paragraph hoặc table trong học liệu (HL) - CHỈ XỬ LÝ NỘI DUNG, KHÔNG XỬ LÝ ALIGNMENT."""
        
            # ✅ MỞ RỘNG: hỗ trợ cả Table
            if isinstance(p, DocxTable):
                return self.convert_table_to_html(p, is_hoc_lieu=True)

            # Nếu không phải Paragraph hoặc Table → trả về rỗng
            if not isinstance(p, Paragraph):
                print(f"[WARN] convert_paragraph_for_hl nhận đầu vào không hợp lệ: {type(p)}")
                return "<br>"

            try:
                # 1. CẮT 'HL:' nếu có
                full_text = p.text
                hl_match = re.match(r"^\s*(H\s*L\s*[:：\-]\s*)", full_text, re.IGNORECASE)
                hl_cut_pos = hl_match.end() if hl_match else 0
                
                # 2. XÂY DỰNG HTML từ runs (sau khi cắt HL:)
                html = ""
                current_pos = 0
                
                for run in p.runs:
                    run_text = run.text or ""
                    if not run_text:
                        continue
                    
                    run_start = current_pos
                    run_end = current_pos + len(run_text)
                    current_pos = run_end
                    
                    # Bỏ qua phần text nằm trong vùng HL:
                    if run_end <= hl_cut_pos:
                        continue
                    
                    if run_start < hl_cut_pos:
                        offset = hl_cut_pos - run_start
                        effective_text = run_text[offset:]
                    else:
                        effective_text = run_text
                    
                    if not effective_text:
                        continue
                    
                    seg = self.escape_html(effective_text)
                    
                    # Áp dụng format
                    if run.bold:
                        seg = f"<strong>{seg}</strong>"
                    if run.italic:
                        seg = f"<i>{seg}</i>"
                    if run.underline:
                        seg = f"<u>{seg}</u>"
                    if getattr(run.font, 'superscript', False):
                        seg = f"<sup>{seg}</sup>"
                    if getattr(run.font, 'subscript', False):
                        seg = f"<sub>{seg}</sub>"
                    if getattr(run.font, 'strike', False) or getattr(run, 'strike', False):
                        seg = f"<strike>{seg}</strike>"
                    
                    html += seg

                # 3. XỬ LÝ ẢNH từ runs
                for run in p.runs:
                    try:
                        imgs = self._get_image_tags_from_run(run)
                        if imgs:
                            html += "".join(imgs)
                    except Exception as e:
                        print(f"[WARN] Lỗi _get_image_tags_from_run trong run: {e}")

                # 4. XỬ LÝ ẢNH DRAWING TRỰC TIẾP
                try:
                    drawings = p._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    for drawing in drawings:
                        blip = drawing.find('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                        if blip is not None:
                            rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if rId:
                                width_emu, height_emu = self.lay_kich_thuoc_tu_word_xml(drawing)
                                img_tag = self._make_img_tag_from_rid(rId, width_emu, height_emu)
                                if img_tag:
                                    html += img_tag
                except Exception as e:
                    print(f"[ERROR] Lỗi xử lý drawing trực tiếp: {e}")
                    import traceback
                    traceback.print_exc()

                # 5. ÁP DỤNG THỤT LỀ (KHÔNG XỬ LÝ ALIGNMENT Ở ĐÂY)
                html = html.strip()
                if not html:
                    return "<br>"

                # Thêm thụt lề trái
                leading_spaces = self.get_indent_html(p)
                html = leading_spaces + html

                # CHỈ TRẢ VỀ NỘI DUNG + <br/>, KHÔNG XỬ LÝ ALIGNMENT
                return html + "<br>"

            except Exception as e:
                print(f"[ERROR] convert_paragraph_for_hl: {e}")
                import traceback
                traceback.print_exc()
                return ""




    # def convert_table_to_html(self, table: DocxTable, is_hoc_lieu=False) -> str:
    #     print("[DEBUG][convert_table_to_html] === BẮT ĐẦU XỬ LÝ TABLE ===")
    #     html = "<table class='table-material-question'>"
    #     try:
    #         for r_idx, row in enumerate(table.rows):
    #             html += "<tr>"
    #             for c_idx, cell in enumerate(row.cells):
    #                 parts = []
    #                 # Nested tables
    #                 if hasattr(cell, "tables") and cell.tables:
    #                     for nested in cell.tables:
    #                         parts.append(self.convert_table_to_html(nested, is_hoc_lieu))
    #                 # Paragraphs
    #                 if hasattr(cell, "paragraphs"):
    #                     for p in cell.paragraphs:
    #                         # para_html = ""
    #                         # if is_hoc_lieu:
    #                         #     para_html = self.convert_paragraph_for_hl(p)
    #                         # else:
    #                         #     para_html = self.convert_content_to_html(p)
    #                         # if para_html:
    #                         #     # ✅ KHÔNG bọc para_html trong <p>...</p> trong table!
    #                         #     parts.append(para_html)
    #                         if is_hoc_lieu:
    #                             para_html = self.convert_paragraph_for_hl(p)
    #                             if para_html:
    #                                 parts.append(para_html)  # ✅ KHÔNG bọc <p>...</p>
    #                         else:
    #                             para_html = self.convert_content_to_html(p)
    #                             parts.append(para_html)
    #                 cell_html = "".join(parts).strip()
    #                 if not cell_html:
    #                     cell_html = "&nbsp;"
    #                 html += f"<td>{cell_html}</td>"
    #             html += "</tr>"
    #     except Exception as e:
    #         print(f"[ERROR] convert_table_to_html: {e}")
    #         traceback.print_exc()
    #     html += "</table><br>"
    #     return html


    def get_vmerge_value(self, tc_pr):
        """Trích xuất giá trị vMerge, mặc định là 'continue' nếu có thẻ nhưng không có w:val."""
        if tc_pr is None:
            return None
        vmerge = tc_pr.find(qn('w:vMerge'))
        if vmerge is None:
            return None
        val = vmerge.get(qn('w:val'))
        return val if val is not None else 'continue'

    def convert_table_to_html(self, table: DocxTable, is_hoc_lieu=False) -> str:
        # Thêm border, cellpadding, cellspacing như HTML "đúng"
        html = '<table class="table-material-question">'
        grid = []  # grid[r][c] = dict (ô gốc) hoặc "OCCUPIED"

        try:
            rows = table.rows
            n_rows = len(rows)

            # Giai đoạn 1: Phân tích từng dòng với con trỏ cột logic
            for r_idx in range(n_rows):
                row = rows[r_idx]
                while len(grid) <= r_idx:
                    grid.append([])

                logical_col = 0  # Con trỏ cột logic, bắt đầu từ 0 mỗi dòng

                for cell_xml in row._element:
                    if cell_xml.tag != qn('w:tc'):
                        continue

                    tc_pr = cell_xml.find(qn('w:tcPr'))

                    # --- COLSPAN ---
                    grid_span = tc_pr.find(qn('w:gridSpan')) if tc_pr is not None else None
                    colspan = int(grid_span.get(qn('w:val'))) if grid_span is not None else 1

                    # --- Kiểm tra vMerge ---
                    vmerge_val = self.get_vmerge_value(tc_pr)

                    if vmerge_val == "continue":
                        # Đánh dấu các ô bị chiếm trong grid
                        for dc in range(colspan):
                            c = logical_col + dc
                            while len(grid[r_idx]) <= c:
                                grid[r_idx].append(None)
                            grid[r_idx][c] = "OCCUPIED"
                        logical_col += colspan
                        continue

                    # --- Tính ROWSPAN bằng cách dò xuống dưới ---
                    rowspan = 1
                    for rr in range(r_idx + 1, n_rows):
                        next_row = rows[rr]
                        next_logical_col = 0
                        found = False

                        for next_cell in next_row._element:
                            if next_cell.tag != qn('w:tc'):
                                continue

                            next_tc_pr = next_cell.find(qn('w:tcPr'))
                            next_grid_span = next_tc_pr.find(qn('w:gridSpan')) if next_tc_pr is not None else None
                            next_colspan = int(next_grid_span.get(qn('w:val'))) if next_grid_span is not None else 1

                            # Nếu đúng cột logic cần kiểm tra
                            if next_logical_col == logical_col:
                                next_vmerge = self.get_vmerge_value(next_tc_pr)
                                if next_vmerge == "continue":
                                    rowspan += 1
                                    found = True
                                break

                            next_logical_col += next_colspan

                        if not found:
                            break

                    # Tạo cell object
                    cell_obj = _Cell(cell_xml, row)
                    cell_data = {
                        "cell": cell_obj,
                        "xml": cell_xml,
                        "rowspan": rowspan,
                        "colspan": colspan,
                    }

                    # Đánh dấu vào grid
                    for dr in range(rowspan):
                        tr = r_idx + dr
                        while len(grid) <= tr:
                            grid.append([])
                        for dc in range(colspan):
                            tc = logical_col + dc
                            while len(grid[tr]) <= tc:
                                grid[tr].append(None)
                            if dr == 0 and dc == 0:
                                grid[tr][tc] = cell_data
                            else:
                                grid[tr][tc] = "OCCUPIED"

                    logical_col += colspan

            # Giai đoạn 2: Render HTML từ grid
            for row in grid:
                html += "<tr>"
                for cell in row:
                    if not isinstance(cell, dict):
                        continue
                    parts = []
                    for child in cell["xml"]:
                        if child.tag == qn("w:tbl"):
                            nested = DocxTable(child, cell["cell"])
                            parts.append(self.convert_table_to_html(nested, is_hoc_lieu))
                        elif child.tag == qn("w:p"):
                            p = Paragraph(child, cell["cell"])
                            content = (
                                self.convert_paragraph_for_hl(p) if is_hoc_lieu
                                else self.convert_content_to_html(p)
                            )
                            parts.append(content)
                    content = "".join(parts).strip() or "&nbsp;"
                    attrs = []
                    if cell["rowspan"] > 1:
                        attrs.append(f'rowspan="{cell["rowspan"]}"')
                    if cell["colspan"] > 1:
                        attrs.append(f'colspan="{cell["colspan"]}"')
                    html += f"<td {' '.join(attrs)}>{content}</td>"
                html += "</tr>"

        except Exception as e:
            import traceback
            print("[ERROR] convert_table_to_html:", e)
            traceback.print_exc()

        html += "</table>"
        return html

    def wrap_style(self, text, style):
        """Đóng gói text với style tuple"""
        bold, italic, underline, sup, sub, strike = style

        if bold:

            text = f"<strong>{text}</strong>"

        if italic:

            text = f"<i>{text}</i>"

        if underline:

            text = f"<u>{text}</u>"

        if sup:

            text = f"<sup>{text}</sup>"    

        if sub:

            text = f"<sub>{text}</sub>"

        if strike:

            text = f"<strike>{text}</strike>"    

        return text

    def format_questions(self, group, questions_xml, errors):
        """Format các câu hỏi, nhận thêm danh sách errors để ghi lỗi"""
        group_of_q = []
        for para in group['questions']:
            if isinstance(para, Table):
                if group_of_q and group_of_q[-1]:
                    group_of_q[-1]['items'].append(para)
                continue
            text = para.text.strip().lower()
            # Phát hiện câu hỏi mới
            if re.match(r'^c[ââ]u.\d', text):
                question_tag = getattr(para, 'current_tag', None) or group.get('original_tag') or group['tag']
                question = {
                    'items': [para],
                    'question_tag': question_tag
                }
                group_of_q.append(question)
            elif group_of_q:
                group_of_q[-1]['items'].append(para)

        # Xử lý từng câu hỏi
        for idx, question_dict in enumerate(group_of_q):
            each_question_xml = Element('question')
            # Metadata
            SubElement(each_question_xml, 'indexGroupQuestionMaterial').text = str(self.index_question)

            SubElement(each_question_xml, 'subject').text = group['subject']

            question_tag = question_dict['question_tag']

            SubElement(each_question_xml, 'tag').text = question_tag

            SubElement(each_question_xml, 'posttype').text = group['posttype']

            SubElement(each_question_xml, 'knowledgelevel').text = group['knowledgelevel']

            SubElement(each_question_xml, 'levelquestion').text = str(group['level'])
            # Xử lý nội dung câu hỏi
            try:
                # Gọi protocol_of_q với danh sách lỗi
                self.protocol_of_q(question_dict['items'], each_question_xml, group['subject'], errors, idx + 1) # idx+1 là số thứ tự câu hỏi
            except Exception as e:
                # Nếu protocol_of_q ném lỗi không bắt được (nên ít xảy ra sau khi sửa)
                # thì vẫn ghi vào danh sách lỗi và tiếp tục
                error_msg = f"Lỗi không xử lý được khi phân tích câu hỏi {idx + 1}: {str(e)}"
                errors.append(error_msg)
                print(f"[ERROR] format_questions: {error_msg}")
                traceback.print_exc()
                continue # Bỏ qua câu hỏi lỗi, tiếp tục với câu tiếp theo

            self.index_question += 1
            questions_xml.append(each_question_xml)

   

    def _get_image_tags_from_run(self, run):
        """
        Trích xuất ảnh từ run, tính KÍCH THƯỚC HIỂN THỊ theo chuẩn Google Docs (pixel GAS).
        Không dùng Pillow.size để xác định kích thước hiển thị.
        """
        imgs = []
        try:
            r = run._r
            from lxml import etree
            nsmap = {
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                'v': 'urn:schemas-microsoft-com:vml'
            }

            # --- 1. DrawingML: blip + extent ---
            blips = r.findall('.//a:blip', nsmap)
            extents = r.findall('.//wp:extent', nsmap)

            for idx, blip in enumerate(blips):
                rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if not rId:
                    continue

                display_width_px, display_height_px = None, None

                # Dùng extent để tính KÍCH THƯỚC HIỂN THỊ (pixel trong Google Docs)
                if idx < len(extents):
                    extent = extents[idx]
                    cx = extent.get('cx')  # EMU
                    cy = extent.get('cy')  # EMU
                    if cx and cy:
                        # Google Docs dùng DPI ≈ 220 cho hiển thị
                        # 1 inch = 220 pixel (GAS), 1 inch = 914400 EMU
                        # → 1 EMU = 220 / 914400 pixel
                        display_width_px = int(int(cx) * 220 / 914400)
                        display_height_px = int(int(cy) * 220 / 914400)

                img_tag = self._make_img_tag_from_rid(rId, display_width_px, display_height_px)
                if img_tag:
                    imgs.append(img_tag)

            # --- 2. VML (hiếm, nhưng xử lý nếu có) ---
            picts = r.findall('.//v:imagedata', nsmap)
            for pict in picts:
                rId = pict.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                if not rId:
                    continue

                display_width_px, display_height_px = None, None
                shape = pict.getparent()
                if shape is not None:
                    style = shape.get('style', '')
                    import re
                    width_match = re.search(r'width:\s*(\d+(?:\.\d+)?)pt', style)
                    height_match = re.search(r'height:\s*(\d+(?:\.\d+)?)pt', style)
                    if width_match and height_match:
                        # Chuyển pt → inch → pixel (220 DPI)
                        width_pt = float(width_match.group(1))
                        height_pt = float(height_match.group(1))
                        # 1 pt = 1/72 inch → pixel = (pt / 72) * 220
                        display_width_px = int(width_pt * 220 / 72)
                        display_height_px = int(height_pt * 220 / 72)

                img_tag = self._make_img_tag_from_rid(rId, display_width_px, display_height_px)
                if img_tag:
                    imgs.append(img_tag)

        except Exception as e:
            print(f"[ERROR] _get_image_tags_from_run: {e}")
            import traceback
            traceback.print_exc()
        return imgs


    def lay_kich_thuoc_tu_word_xml(self,drawing_element):
        """
        Lấy cx, cy (EMU units) từ Word XML.
        
        Ví dụ XML:
        <wp:extent cx="2006920" cy="1828800"/>
        
        Returns:
            (width_emu, height_emu) hoặc (None, None)
        """
        try:
            # Namespace của Word XML
            namespaces = {
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
            }
            
            # Tìm extent element
            extent = drawing_element.find('.//wp:extent', namespaces)
            
            if extent is not None:
                cx_emu = int(extent.get('cx', 0))  # width in EMU
                cy_emu = int(extent.get('cy', 0))  # height in EMU
                
                print(f"[DEBUG] ✓ Tìm thấy extent: cx={cx_emu} EMU, cy={cy_emu} EMU")
                return cx_emu, cy_emu
            else:
                print(f"[DEBUG] ✗ Không tìm thấy wp:extent trong drawing")
                return None, None
                
        except Exception as e:
            print(f"[ERROR] lay_kich_thuoc_tu_word_xml: {e}")
            import traceback
            traceback.print_exc()
            return None, None

    def _make_img_tag_from_rid(self, rId, display_width_emu=None, display_height_emu=None):
        print(f">>>>>>>> chiều rộng emu {display_width_emu}")

        print(f">>>>>>>>> chiều dài emu {display_height_emu}")
        """
        Tạo thẻ <img> với kích thước CHÍNH XÁC từ Word XML.
        
        QUAN TRỌNG:
        - Phải truyền display_width_emu và display_height_emu từ Word XML
        - ĐỪNG dùng img.size (pixel vật lý của ảnh gốc)
        - Google Apps Script dùng EMU units từ Word XML
        
        Công thức:
        - Word lưu: cx, cy (EMU units)
        - GAS getWidth() = cx / 12700 (point)
        - HTML style: width = (cx / 12700)px
        """
        try:
            part = self.doc.part.related_parts.get(rId)
            if not part:
                for rel in self.doc.part.rels.values():
                    try:
                        target = getattr(rel, 'target_part', None)

                        if target and 'image' in getattr(target, 'content_type', ''):

                            if rel.rId == rId:

                                part = target

                                break
                    except Exception:
                        continue

            if not part:
                print(f"[DEBUG] Không tìm thấy part cho rId={rId}")
                return None

            img_bytes = part.blob

            content_type = getattr(part, 'content_type', 'image/png')
            
            # === TÍNH KÍCH THƯỚC TỪ WORD XML EMU ===
            if display_width_emu is not None and display_height_emu is not None:
                # Chuyển EMU → point (khớp với GAS)
                # 1 point = 12700 EMU
                final_width = round(display_width_emu / 9525)


                final_height = round(display_height_emu / 9525)

                print(f"[DEBUG] Word XML: {display_width_emu}x{display_height_emu} EMU")

                print(f"[DEBUG] GAS output: {final_width}x{final_height} pt")
            else:
                # FALLBACK: Dùng kích thước ảnh gốc (KHÔNG KHUYẾN NGHỊ)
                img = Image.open(BytesIO(img_bytes))

                pixel_width, pixel_height = img.size

                dpi_info = img.info.get('dpi', (96, 96))

                dpi = dpi_info[0] if isinstance(dpi_info, tuple) else dpi_info
                
                final_width = round(pixel_width * 72 / dpi)

                final_height = round(pixel_height * 72 / dpi)

                print(f"[WARNING] Không có EMU từ Word XML, dùng fallback!")

                print(f"[DEBUG] Ảnh gốc: {pixel_width}x{pixel_height} px @ {dpi} DPI")

                print(f"[DEBUG] Fallback: {final_width}x{final_height} pt")

            # KHÔNG RESIZE - giữ nguyên ảnh gốc
            output = BytesIO()

            img = Image.open(BytesIO(img_bytes))

            img_format = img.format or 'PNG'

            img.save(output, format=img_format, optimize=False)

            b64 = base64.b64encode(output.getvalue()).decode('ascii')

            output.close()

            return f'<center><img style="width:{final_width}px; height:{final_height}px;" src="data:{content_type};base64,{b64}" /></center>'

        except Exception as e:
            print(f"[ERROR] _make_img_tag_from_rid: {e}")
            import traceback
            traceback.print_exc()
            return None
        
    def get_hyperlinks_from_paragraph(self,paragraph: Paragraph):
        links = []
        part = paragraph.part

        for hyperlink in paragraph._p.findall(qn('w:hyperlink')):
            r_id = hyperlink.get(qn('r:id'))
            if r_id:
                url = part.rels[r_id].target_ref
                links.append(url)

        return links

    # def protocol_of_q(self, question, each_question_xml, subject, errors, question_index):
    #     """Phân tích cấu trúc câu hỏi, nhận danh sách errors và số thứ tự câu hỏi question_index"""
    #     # Chia thành phần: nội dung câu hỏi và lời giải
    #     thanh_phan_1q = []

    #     for idx, para in enumerate(question):

    #         if idx == 0:

    #             thanh_phan_1q.append([para])

    #             continue
    #         if isinstance(para, Paragraph):

    #             text = para.text.strip().lower()
    #             # print(f">>>>>> debug text phan loai: {text}")

    #             # if re.match(r'^l[ờờ]i gi[ảả]i', text):
    #             if re.match(r'^\s*l[ờơ]i\s+gi[ảẩ]i\s*[:：]?', text, re.IGNORECASE):

    #                 thanh_phan_1q.append([])

    #                 continue
    #         if thanh_phan_1q:
    #             thanh_phan_1q[-1].append(para)

    #     if len(thanh_phan_1q) < 2:
    #         # raise ValueError(f"Thiếu 'Lời giải' trong câu: {question[0].text[:50]}")
    #         error_msg = f"Thiếu 'Lời giải' trong câu hỏi {question_index}"

    #         errors.append(error_msg)

    #         print(f"[ERROR] protocol_of_q: {error_msg}")

    #         SubElement(each_question_xml, 'contentquestion').text = ''

    #         SubElement(each_question_xml, 'explainquestion').text = f'--- LỖI: Thiếu lời giải ---'

    #         SubElement(each_question_xml, 'typeAnswer').text = '0' # Mặc định
            
    #         return # Kết thúc xử lý câu hỏi này

    #     # Phân tích nội dung câu hỏi và lời giải
    #     thanh_phan_cau_hoi = []

    #     link_cau_hoi = []

    #     for idx, para in enumerate(thanh_phan_1q[0]):
    #         if isinstance(para, Paragraph):
    #             text = para.text.strip()
    #             print(f">>>>>> debug text cau hoi: {text}")
    #             hyperlinks = self.get_hyperlinks_from_paragraph(para)

    #             for link in hyperlinks:
    #                 if link not in link_cau_hoi:
    #                     link_cau_hoi.append(link)
    #             # ——— XỬ LÝ DÒNG BẮT ĐẦU BẰNG "Audio:" ———
    #             if text.startswith('Audio:'):
    #                 # print(f">>>>>> debug audio content: {audio_content}")
    #                 audio_content = text[6:].strip()
                   
    #                 # Nếu ngay sau có link hợp lệ → dùng luôn
    #                 if audio_content.startswith('https'):
    #                     link_cau_hoi.append(f'Audio:{audio_content}')
    #                 else:
    #                     # Nếu không, kiểm tra paragraph tiếp theo có URL không
    #                     if idx + 1 < len(thanh_phan_1q[0]):

    #                         next_para = thanh_phan_1q[0][idx + 1]

    #                         if isinstance(next_para, Paragraph):

    #                             next_text = next_para.text.strip()

    #                             # Kiểm tra link thuần hoặc link có hyperlink (giả lập: chỉ kiểm tra text)

    #                             if next_text.startswith('https'):

    #                                 link_cau_hoi.append(f'Audio:{next_text}')

    #                                 # Bỏ qua para tiếp theo trong nội dung chính
    #                                 # (nhưng vẫn giữ nguyên logic append → sẽ loại sau)
    #                 continue  # Dù thế nào cũng không đưa "Audio:" vào nội dung chính

    #             # ——— XỬ LÝ URL THUẦN TRONG ĐOẠN VĂN ———
    #             # Tìm mọi URL hợp lệ trong text (kể cả link bị kèm chữ)
    #             # url_matches = re.findall(r'https?://[^\s]+', text)
    #             # # print(f">>>>>> debug url matches: {url_matches}")
    #             # found_valid_url = False
    #             # for url in url_matches:
    #             #     url_clean = url.rstrip('.,;:')
    #             #     if url_clean not in [link.replace('Audio:', '', 1) for link in link_cau_hoi]:
    #             #         link_cau_hoi.append(url_clean)
    #             #         found_valid_url = True
    #             # # Nếu URL đứng riêng (không kèm text quan trọng), không thêm vào nội dung chính
    #             # if url_matches and not text[:text.find(url_matches[0])].strip():
    #             #     continue

    #             url_matches = re.findall(r'https?://[^\s]+', text)
    #             found_valid_url = False

    #             # Kiểm tra hyperlink trong các run
    #             if isinstance(para, Paragraph):
    #                 for run in para.runs:
    #                     # Kiểm tra hyperlink trong run
    #                     if run._element.rPr is not None:
    #                         rpr = run._element.rPr
    #                         # Tìm hyperlink
    #                         hyperlinks = run._element.xpath('.//w:hyperlink')
    #                         for hyperlink in hyperlinks:
    #                             r_id = hyperlink.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
    #                             if r_id:
    #                                 # Lấy relationship từ document
    #                                 rel = para.part.rels[r_id]
    #                                 if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink':
    #                                     url = rel.target_ref
    #                                     if url.startswith('http'):
    #                                         link_cau_hoi.append(url)
    #                                         found_valid_url = True
    #                                         print(f">>>>>> [HYPERLINK DETECTED] {url}")

    #             # Thêm URL thuần (nếu có)
    #             for url in url_matches:
    #                 url_clean = url.rstrip('.,;:')
    #                 if url_clean not in [link.replace('Audio:', '', 1) for link in link_cau_hoi]:
    #                     link_cau_hoi.append(url_clean)
    #                     found_valid_url = True
    #                     print(f">>>>>> [PLAIN URL DETECTED] {url_clean}")
    #             # Thêm vào nội dung chính nếu không phải dòng audio hoặc link thuần
    #             thanh_phan_cau_hoi.append(para)
    #     # Xử lý links
    #     self.xu_ly_link_cau_hoi(link_cau_hoi, each_question_xml)

    #     # Phân tích lời giải
    #     thanh_phan_hdg = []

    #     link_speech_explain = []

    #     has_sharpened = False

    #     for idx, para in enumerate(thanh_phan_1q[1]):
    #         if idx == 0:

    #             thanh_phan_hdg.append([para])
    #             continue

    #         if isinstance(para, Paragraph):

    #             text = para.text.strip()
    #             print(f">>>>>> debug text loi giai: {text}")
               

    #             if text.startswith('###'):
    #                 has_sharpened = True
    #                 thanh_phan_hdg.append([])

    #                 continue
    #             # URLs trong HDG
    #             urls = re.findall(r'http?://[^\s]+', text)

    #             for url in urls:

    #                 link_speech_explain.append(url)

    #                 continue

    #         if thanh_phan_hdg:

    #             thanh_phan_hdg[-1].append(para)

    #     # Xử lý urlSpeechExplain
    #     if link_speech_explain:

    #         for link in link_speech_explain:
    #             if link.endswith(('.mp3', '.mp4')):
    #                 SubElement(each_question_xml, 'urlSpeechExplain').text = link

    #     # Xác định dạng câu hỏi
    #     answer = thanh_phan_hdg[0][0].text.strip() if thanh_phan_hdg[0] else ''

    #     cau_sau_xu_ly = [thanh_phan_cau_hoi, thanh_phan_hdg]

    #     # audio = [link for link in link_cau_hoi if 'Audio:' in link]
    #     audio = []

    #     for item in question:

    #         if isinstance(item, Paragraph):

    #             txt = item.text.strip()
    #             if txt.startswith('Audio:'):
    #                 print(f">>>>>> debug txt have audio {txt}")

    #                 audio.append(txt)
    #             # if txt.startswith('https://mathplay.onluyen.vn'):
    #             #     print(f">>>>>> debug txt have audio {txt}")

    #                 audio.append(txt)

    #     print(f">>>>>>>>> debug has_sharpened: {has_sharpened}")

    #     # Routing theo subject
    #     if self.is_tinhoc_subject(subject):
    #         self.route_to_tinhoc_module(cau_sau_xu_ly, each_question_xml, audio, answer, subject, errors, question_index)
    #     else:
    #         self.route_to_default_module(cau_sau_xu_ly, each_question_xml, audio, answer, subject, errors, question_index,has_sharpened)
   

    def protocol_of_q(self, question, each_question_xml, subject, errors, question_index):
        """Phân tích cấu trúc câu hỏi, nhận danh sách errors và số thứ tự câu hỏi question_index"""
        # Chia thành phần: nội dung câu hỏi và lời giải
        thanh_phan_1q = []

        for idx, para in enumerate(question):
            if idx == 0:
                thanh_phan_1q.append([para])
                continue
            if isinstance(para, Paragraph):
                text = para.text.strip().lower()
                if re.match(r'^\s*l[ờơ]i\s+gi[ảẩ]i\s*[:：]?', text, re.IGNORECASE):
                    thanh_phan_1q.append([])
                    continue
            if thanh_phan_1q:
                thanh_phan_1q[-1].append(para)

        if len(thanh_phan_1q) < 2:
            error_msg = f"Thiếu 'Lời giải' trong câu hỏi {question_index}"
            errors.append(error_msg)
            print(f"[ERROR] protocol_of_q: {error_msg}")
            SubElement(each_question_xml, 'contentquestion').text = ''
            SubElement(each_question_xml, 'explainquestion').text = f'--- LỖI: Thiếu lời giải ---'
            SubElement(each_question_xml, 'typeAnswer').text = '0'
            return

        # Phân tích nội dung câu hỏi và lời giải
        thanh_phan_cau_hoi = []
        link_cau_hoi = []

        for idx, para in enumerate(thanh_phan_1q[0]):
            if isinstance(para, Paragraph):
                text = para.text.strip()
                print(f">>>>>> debug text cau hoi: {text}")
                
                # ===== FIX: DETECT HYPERLINK TRƯỚC TIÊN =====
                # 1. Lấy hyperlink từ paragraph (method có sẵn)
                hyperlinks = self.get_hyperlinks_from_paragraph(para)
                for link in hyperlinks:
                    if link not in link_cau_hoi:
                        link_cau_hoi.append(link)
                        print(f">>>>>> [HYPERLINK VIA METHOD] {link}")
                
                # 2. Detect hyperlink trực tiếp từ XML structure
                for run in para.runs:
                    # Tìm hyperlink element trong run
                    hyperlink_elements = run._element.xpath('.//w:hyperlink', 
                        namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    
                    for hyperlink_elem in hyperlink_elements:
                        r_id = hyperlink_elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                        if r_id and r_id in para.part.rels:
                            rel = para.part.rels[r_id]
                            if rel.reltype == 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink':
                                url = rel.target_ref
                                if url and url.startswith('http'):
                                    if url not in link_cau_hoi:
                                        link_cau_hoi.append(url)
                                        print(f">>>>>> [HYPERLINK VIA XML] {url}")

                # ===== XỬ LÝ DÒNG "Audio:" =====
                if text.startswith('Audio:'):
                    audio_content = text[6:].strip()
                    
                    # Nếu ngay sau có link hợp lệ → dùng luôn
                    if audio_content.startswith('http'):
                        if f'Audio:{audio_content}' not in link_cau_hoi:
                            link_cau_hoi.append(f'Audio:{audio_content}')
                    else:
                        # Kiểm tra paragraph tiếp theo
                        if idx + 1 < len(thanh_phan_1q[0]):
                            next_para = thanh_phan_1q[0][idx + 1]
                            if isinstance(next_para, Paragraph):
                                next_text = next_para.text.strip()
                                if next_text.startswith('http'):
                                    if f'Audio:{next_text}' not in link_cau_hoi:
                                        link_cau_hoi.append(f'Audio:{next_text}')
                    continue  # Không đưa dòng Audio: vào nội dung chính

                # ===== XỬ LÝ URL THUẦN (plain text URLs) =====
                url_matches = re.findall(r'https?://[^\s]+', text)
                is_url_only_para = False
                
                if url_matches:
                    # Kiểm tra xem paragraph có phải chỉ chứa URL không
                    text_without_urls = text
                    for url in url_matches:
                        text_without_urls = text_without_urls.replace(url, '')
                    text_without_urls = text_without_urls.strip()
                    
                    # Nếu sau khi bỏ URL, không còn nội dung quan trọng
                    is_url_only_para = len(text_without_urls) == 0
                    
                    # Thêm các URL vào danh sách
                    for url in url_matches:
                        url_clean = url.rstrip('.,;:')
                        # Tránh duplicate với Audio: prefix
                        already_exists = any(
                            link == url_clean or link == f'Audio:{url_clean}' 
                            for link in link_cau_hoi
                        )
                        if not already_exists:
                            link_cau_hoi.append(url_clean)
                            print(f">>>>>> [PLAIN URL] {url_clean}")
                    
                    # Nếu paragraph chỉ chứa URL, không thêm vào nội dung
                    if is_url_only_para:
                        continue

                # Thêm vào nội dung câu hỏi (nếu không phải Audio: hoặc URL thuần)
                thanh_phan_cau_hoi.append(para)

        # Xử lý links
        self.xu_ly_link_cau_hoi(link_cau_hoi, each_question_xml)

        # Phân tích lời giải
        thanh_phan_hdg = []
        link_speech_explain = []
        has_sharpened = False

        for idx, para in enumerate(thanh_phan_1q[1]):
            if idx == 0:
                thanh_phan_hdg.append([para])
                continue

            if isinstance(para, Paragraph):
                text = para.text.strip()
                print(f">>>>>> debug text loi giai: {text}")

                if text.startswith('###'):
                    has_sharpened = True
                    thanh_phan_hdg.append([])
                    continue
                
                # URLs trong HDG
                urls = re.findall(r'https?://[^\s]+', text)
                for url in urls:
                    link_speech_explain.append(url)
                    continue

            if thanh_phan_hdg:
                thanh_phan_hdg[-1].append(para)

        # Xử lý urlSpeechExplain
        if link_speech_explain:
            for link in link_speech_explain:
                if link.endswith(('.mp3', '.mp4')):
                    SubElement(each_question_xml, 'urlSpeechExplain').text = link

        # Xác định dạng câu hỏi
        answer = thanh_phan_hdg[0][0].text.strip() if thanh_phan_hdg[0] else ''
        cau_sau_xu_ly = [thanh_phan_cau_hoi, thanh_phan_hdg]

        # Detect audio từ question list
        audio = []
        for item in question:
            if isinstance(item, Paragraph):
                txt = item.text.strip()
                if txt.startswith('Audio:') or txt.startswith('https://mathplay.onluyen.vn'):
                    print(f">>>>>> debug txt have audio {txt}")
                    audio.append(txt)

        print(f">>>>>>>>> debug has_sharpened: {has_sharpened}")

        # Routing theo subject
        if self.is_tinhoc_subject(subject):
            self.route_to_tinhoc_module(cau_sau_xu_ly, each_question_xml, audio, answer, subject, errors, question_index)
        else:
            self.route_to_default_module(cau_sau_xu_ly, each_question_xml, audio, answer, subject, errors, question_index, has_sharpened)




    def is_tinhoc_subject(self, subject):
        """Kiểm tra có phải môn tin học không"""
        return any(subject.startswith(tinhoc) for tinhoc in self.tinhoc_subjects)

    def route_to_tinhoc_module(self, cau_sau_xu_ly, xml, audio, answer, subject, errors, question_index):
        """Xử lý cho môn Tin học, nhận danh sách lỗi và số câu hỏi"""
        # ✅ Gọi từ instance tinhoc_processor
        if re.match(r'^\d+', answer):
            if len(answer) > 1 and re.match(r'^[01]+', answer):
                self.tinhoc_processor.dang_ds_tinhoc(cau_sau_xu_ly, xml, audio, self.doc)
            else:
                self.tinhoc_processor.dang_tn_tinhoc(cau_sau_xu_ly, xml, audio, self.doc)
        elif answer.startswith('##'):
            self.dang_dt(cau_sau_xu_ly, xml, subject)
        else:
            self.dang_tl(cau_sau_xu_ly, xml, audio)

    def route_to_default_module(self, cau_sau_xu_ly, xml, audio, answer, subject, errors, question_index,has_sharpened):
        """Xử lý cho môn thông thường, nhận danh sách lỗi và số câu hỏi"""
        if re.match(r'^\d+', answer):
            if len(answer) > 1 and re.match(r'^[01]+', answer):
                print(f">>>>>  Default → Dang Dung/Sai")
                self.dang_ds(cau_sau_xu_ly, xml, audio)
            else:
                print(f">>>>>  Default → Dang Trac Nghiem")
                self.dang_tn(cau_sau_xu_ly, xml, audio)
        elif answer.startswith('##'):
            print(f">>>>>  Default → Dang Dien Tu")
            self.dang_dt(cau_sau_xu_ly, xml, subject)
        else:
            print(f">>>>>  Default → Dang Tu Luan")
            self.dang_tl(cau_sau_xu_ly, xml, audio)

        """
        Xử lý cho môn thông thường, nhận danh sách lỗi và số câu hỏi.
        
        Logic theo GAS:
        - Chỉ xử lý TN/DS khi: has_sharpened === True VÀ answer là số thuần
        - Xử lý Điền Từ khi: answer bắt đầu bằng ##
        - Các trường hợp còn lại: Tự luận
        """
        
        # ===== ĐIỀU KIỆN 1: has_sharpened === True VÀ answer là số thuần =====
      
        # answer = answer.strip()

        # # ===== 1. ĐIỀN TỪ =====
        # if answer.startswith('##'):
        #     print('Default → Dang Dien Tu')
        #     self.dang_dt(cau_sau_xu_ly, xml, subject)
        #     return

        # # ===== 2. TRẮC NGHIỆM / ĐÚNG SAI =====
        # if has_sharpened is True and re.fullmatch(r'\d+', answer):

        #     # ĐÚNG / SAI: 1010, 0110
        #     if len(answer) > 1 and re.fullmatch(r'[01]+', answer):
        #         print('Default → Dang Dung/Sai')
        #         self.dang_ds(cau_sau_xu_ly, xml, audio)
        #         return

        #     # TRẮC NGHIỆM: 1,2,3,4
        #     print('Default → Dang Trac Nghiem')
        #     self.dang_tn(cau_sau_xu_ly, xml, audio)
        #     return

        # # ===== 3. TỰ LUẬN =====
        # print('Default → Dang Tu Luan')
        # self.dang_tl(cau_sau_xu_ly, xml, audio)

    # def xu_ly_link_cau_hoi(self, links: str, xml):
    #     """Xử lý links trong câu hỏi"""
    #     one_tts = False

    #     one_media = False

    #     for link in links:

    #         if link.startswith('Audio:'):

    #             continue

    #         if link.endswith(('.mp3', '.mp4')):

    #             if one_tts:

    #                 # raise ValueError(f"Chỉ được 1 link TTS: {link}")
    #                 print(f"[WARN] Có nhiều hơn 1 link TTS trong câu hỏi, bỏ qua: {link}")

    #                 continue

    #             SubElement(xml, 'urlSpeechContent').text = link

    #             one_tts = True
    #         else:
    #             if one_media:

    #                 # raise ValueError(f"Chỉ được 1 link Video: {link}")
    #                 print(f"[WARN] Có nhiều hơn 1 link Video trong câu hỏi, bỏ qua: {link}")

    #                 continue

    #             if 'vimeo.com' in link:

    #                 code = link.split('vimeo.com/')[1]

    #                 parts = code.split('/')

    #                 if len(parts) > 1:

    #                     code = f"{parts[0]}?h={parts[1].split('?share')[0]}"
    #                 else:
    #                     code = parts[0]

    #                 SubElement(xml, 'contentMedia').text = code

    #                 SubElement(xml, 'typeContentMedia').text = 'CodeVimeo'

    #                 one_media = True

    #             elif 'youtu' in link:

    #                 if 'watch?v=' in link:

    #                     code = link.split('watch?v=')[1]

    #                 elif 'youtu.be/' in link:

    #                     code = link.split('youtu.be/')[1].split('?')[0]

    #                 else:
    #                     continue
    #                 SubElement(xml, 'contentMedia').text = code

    #                 SubElement(xml, 'typeContentMedia').text = 'CodeYouTuBe'
    #                 one_media = True

    def xu_ly_link_cau_hoi(self, links: List[str], xml):
        """Xử lý links trong câu hỏi — ĐÃ CẬP NHẬT LOGIC TTS"""
        one_tts = False
        one_media = False
        for link in links:
            # ——— Chuẩn hóa link ———
            clean_link = link
            if link.startswith('Audio:'):
                clean_link = link[6:].strip()
            else:
                clean_link = link.strip()

            # Bỏ qua nếu rỗng
            if not clean_link:
                continue

            # ——— PHÂN LOẠI LINK ———
            if clean_link.endswith(('.mp3', '.mp4')):
                if one_tts:
                    print(f"[WARN] Có nhiều hơn 1 link TTS trong câu hỏi, bỏ qua: {clean_link}")
                    continue
                SubElement(xml, 'urlSpeechContent').text = clean_link
                one_tts = True
            else:
                if one_media:
                    print(f"[WARN] Có nhiều hơn 1 link Video trong câu hỏi, bỏ qua: {clean_link}")
                    continue
                if 'vimeo.com' in clean_link:
                    code = clean_link.split('vimeo.com/')[1]
                    parts = code.split('/')
                    if len(parts) > 1:
                        code = f"{parts[0]}?h={parts[1].split('?share')[0]}"
                    else:
                        code = parts[0]
                    SubElement(xml, 'contentMedia').text = code
                    SubElement(xml, 'typeContentMedia').text = 'CodeVimeo'
                    one_media = True
                elif 'youtu' in clean_link:
                    if 'watch?v=' in clean_link:
                        code = clean_link.split('watch?v=')[1]
                    elif 'youtu.be/' in clean_link:
                        code = clean_link.split('youtu.be/')[1].split('?')[0]
                    else:
                        continue
                    SubElement(xml, 'contentMedia').text = code
                    SubElement(xml, 'typeContentMedia').text = 'CodeYouTuBe'
                    one_media = True

    # ... (các hàm convert_content_to_html, dang_tn, list_answers_tn, strip_html, hdg_tn, dang_ds, dang_dt, dang_tl, convert_b4_add, convert_normal_paras, escape_html, prettify_xml) ...
    # Các hàm này không cần thay đổi để phù hợp với cơ chế mới, trừ khi chúng có thể ném lỗi và cần được xử lý riêng.
    # Tuy nhiên, để an toàn, ta có thể bao bọc các hàm chính được gọi từ format_questions trong try-except.


    def detect_soft_breaks_in_paragraph(self, p: Paragraph):
        """In ra vị trí và số lượng các soft break (Shift+Enter) trong paragraph để debug"""
        from docx.oxml.ns import qn
        br_nodes = p._element.findall(qn('w:br'))
        soft_breaks = [br for br in br_nodes if br.get(qn('w:type')) == 'textWrapping']
        if soft_breaks:
            print(f"[DEBUG] Phát hiện {len(soft_breaks)} soft break (Shift+Enter) trong paragraph: '{p.text[:300]}...'")
            for i, br in enumerate(soft_breaks):
                # In vị trí tương đối (không chính xác tuyệt đối, nhưng đủ để nhận biết)
                parent = br.getparent()
                if parent is not None:
                    idx = list(parent).index(br)
                    print(f"  → Soft break #{i+1} tại vị trí XML index: {idx}")
        else:
            print(f"[DEBUG] Không có soft break trong paragraph: '{p.text[:50]}...'")

    def convert_content_to_html(self, paragraphs):
        """
        Chuyển list Paragraph / Table sang HTML hoàn chỉnh, giữ table, ảnh, math-latex.
        KHÔNG tự bọc <div class='content'> để tránh lặp.
        Hỗ trợ flatten đệ quy: chấp nhận paragraphs là Paragraph, Table,
        list/tuple lồng nhau ở bất kỳ mức độ nào.
        """
        from docx.table import Table
        # Đệ quy flatten: trả về list các phần tử không phải list/tuple nữa
        def _flatten(items):
            for it in items:
                if isinstance(it, (list, tuple)):
                    yield from _flatten(it)
                else:
                    yield it
        # Nếu người gọi chuyền 1 object không phải iterable (ví dụ một Paragraph),
        # ta chuẩn hóa thành list để xử lý thống nhất.
        if paragraphs is None:
            flat = []
        elif isinstance(paragraphs, (list, tuple)):
            flat = list(_flatten(paragraphs))
        else:
            # Một phần tử đơn lẻ (có thể là Paragraph hoặc Table)
            flat = [paragraphs]
        string_content = ""
        for para in flat:
            # Bảo vệ: nếu para là None thì bỏ qua
            if para is None:
                continue
            # Nếu là Table (obj từ python-docx), xử lý riêng
            if isinstance(para, Table):
                string_content += self.convert_table_to_html(para)
                string_content += "<br>"
                continue
            # Nếu là string (đã chuyển trước đó), thêm trực tiếp
            if isinstance(para, str):
                string_content += para + "<br>"
                continue
            # Một số đối tượng paragraph-like có thể không đến từ python-docx
            # nhưng có attribute 'runs' — kiểm tra trước khi gọi convert_normal_paras
            new_children = []
            try:
                # Nếu paragraph không phải object paragraph hợp lệ, convert_normal_paras có thể ném
                self.convert_normal_paras(para, 0, new_children)
                string_content += "".join(new_children)
            except TypeError:
                # Thử gọi convert_normal_paras theo kiểu cũ (nếu hàm được thiết kế trả về string/list)
                try:
                    res = self.convert_normal_paras(para)
                except Exception as e:
                    # Nếu vẫn lỗi, chuyển sang fallback: str(para)
                    string_content += str(para)
                else:
                    if isinstance(res, str):
                        string_content += res
                    elif isinstance(res, list):
                        string_content += "".join(res)
                    else:
                        string_content += str(res)
            except AttributeError:
                # Thường xảy ra khi para là 1 list lồng mà chưa flatten đúng mức
                # Fallback robust: chuyển thành str(para)
                string_content += str(para)
            string_content += "<br>"
        # Xử lý math-latex
        import re
        math_latex = re.compile(r"\$[^$]*\$")
        string_content = math_latex.sub(lambda m: f'<span class="math-tex">{m.group()}</span>', string_content)
        return string_content.strip()

    # def convert_content_to_html(self, paragraphs):
    #     """
    #     Chuyển đổi danh sách Paragraph / Table / string thành HTML.
    #     Xử lý đúng các dòng trống: mỗi paragraph rỗng → thêm 1 <br>.
    #     Nếu 2 paragraph rỗng liên tiếp → <br><br>.
    #     """
    #     from docx.table import Table
    #     from bs4 import BeautifulSoup

    #     # Hàm đệ quy flatten
    #     def _flatten(items):
    #         for it in items:
    #             if isinstance(it, (list, tuple)):
    #                 yield from _flatten(it)
    #             else:
    #                 yield it

    #     # Chuẩn hóa input
    #     if paragraphs is None:
    #         flat = []
    #     elif isinstance(paragraphs, (list, tuple)):
    #         flat = list(_flatten(paragraphs))
    #     else:
    #         flat = [paragraphs]

    #     string_content = ""
    #     prev_was_empty = False

    #     for para in flat:
    #         if para is None:
    #             # Xử lý None như paragraph rỗng
    #             if prev_was_empty:
    #                 string_content += "<br><br>"
    #             else:
    #                 string_content += "<br>"
    #             prev_was_empty = True
    #             continue

    #         # ——— XỬ LÝ TABLE ———
    #         if isinstance(para, Table):
    #             table_html = self.convert_table_to_html(para)
    #             string_content += table_html + "<br>"
    #             prev_was_empty = False
    #             continue

    #         # ——— XỬ LÝ STRING ———
    #         if isinstance(para, str):
    #             clean_str = para.strip()
    #             is_empty = not clean_str or clean_str in ("<br>", "<br/>")
    #             if is_empty:
    #                 if prev_was_empty:
    #                     string_content += "<br><br>"
    #                 else:
    #                     string_content += "<br>"
    #                 prev_was_empty = True
    #             else:
    #                 string_content += para + "<br>"
    #                 prev_was_empty = False
    #             continue

    #         # ——— XỬ LÝ PARAGRAPH ———
    #         if isinstance(para, Paragraph):
    #             new_children = []
    #             try:
    #                 self.convert_normal_paras(para, 0, new_children)
    #                 para_html = "".join(new_children)
    #                 # Dùng BeautifulSoup để lấy plain text (loại bỏ HTML tags)
    #                 plain_text = BeautifulSoup(para_html, "html.parser").get_text().strip()
    #                 is_empty = not plain_text
    #             except Exception as e:
    #                 # Fallback: coi là có nội dung
    #                 para_html = str(para)
    #                 is_empty = False

    #             if is_empty:
    #                 if prev_was_empty:
    #                     string_content += "<br><br>"
    #                 else:
    #                     string_content += "<br>"
    #                 prev_was_empty = True
    #             else:
    #                 string_content += para_html + "<br>"
    #                 prev_was_empty = False
    #         else:
    #             # Fallback cho các loại khác
    #             fallback_str = str(para)
    #             if fallback_str.strip():
    #                 string_content += fallback_str + "<br>"
    #                 prev_was_empty = False
    #             else:
    #                 if prev_was_empty:
    #                     string_content += "<br><br>"
    #                 else:
    #                     string_content += "<br>"
    #                 prev_was_empty = True

    #     # ——— XỬ LÝ MATH LATEX ———
    #     import re
    #     math_latex = re.compile(r"\$[^$]*\$")
    #     string_content = math_latex.sub(
    #         lambda m: f'<span class="math-tex">{m.group()}</span>',
    #         string_content
    #     )

    #     return string_content.strip()

    def dang_tn(self, cau_sau_xu_ly, xml, audio):
        """
        Xử lý dạng Trắc nghiệm (typeAnswer=0, template=0)
        - Đáp án đúng được xác định bằng số 1,2,3,4 trong phần Lời giải (1=A, 2=B, 3=C, 4=D)
        """
        SubElement(xml, 'typeAnswer').text = '0'
        SubElement(xml, 'typeViewContent').text = '0'
        SubElement(xml, 'template').text = '0'
        # ===== 1️⃣ Xử lý phần nội dung câu hỏi =====
        content_part = []

        answers_part = []

        for para in cau_sau_xu_ly[0]:
            if isinstance(para, Paragraph):

                text = para.text.strip()

                # Nhận diện các dòng A. B. C. D.
                if re.match(r'^[A-Z]\.', text):

                    answers_part.append(para)
                else:

                    content_part.append(para)
            elif isinstance(para, Table):
                content_part.append(para)
        # HTML câu hỏi
        content_html = self.convert_content_to_html(content_part)
        if audio and len(audio[0]) > 8:

            link = audio[0].replace('Audio:', '').strip()

            content_html += f'<audio controls=""><source src="{link}" type="audio/mpeg"></audio>'

        SubElement(xml, 'contentquestion').text = content_html.strip()
        # ===== 2️⃣ Tìm đáp án đúng từ phần Lời giải =====
        correct_index = None  # chỉ số 0-based của đáp án đúng
        if len(cau_sau_xu_ly) > 1 and cau_sau_xu_ly[1]:

            # Lấy đoạn đầu tiên của phần lời giải
            first = cau_sau_xu_ly[1][0]

            if isinstance(first, list):

                # Nếu là danh sách Paragraph
                for p in first:

                    if hasattr(p, 'text'):

                        # m = re.search(r'\b([1-4])\b', p.text.strip())

                        m = re.search(r'\b([1-9]|1[0-9]|2[0-6])\b', p.text.strip())

                        if m:

                            correct_index = int(m.group(1)) - 1

                            break
            elif hasattr(first, 'text'):

                # m = re.search(r'\b([1-4])\b', first.text.strip())
                m = re.search(r'\b([1-9]|1[0-9]|2[0-6])\b', first.text.strip())

                if m:

                    correct_index = int(m.group(1)) - 1
        # ===== 3️⃣ Sinh danh sách đáp án =====
        listanswers = SubElement(xml, 'listanswers')
        for i, para in enumerate(answers_part):
            # Bỏ prefix A./B./C./D.
            # text = re.sub(r'^[A-Z]\.\s*', '', para.text.strip())

            # content_html = f'<p>{text}</p>

            content_html = self.convert_content_to_html([para])

            content_html = re.sub(r'^\s*(?:<[^>]*>)*[A-Z]\.\s*(?:<[^>]*>)*', '', content_html, flags=re.IGNORECASE)

            answer_el = SubElement(listanswers, 'answer')

            SubElement(answer_el, 'index').text = str(i)

            SubElement(answer_el, 'content').text = content_html

            SubElement(answer_el, 'isanswer').text = 'TRUE' if i == correct_index else 'FALSE'
        # ===== 4️⃣ Gọi hdg_tn() để xử lý phần giải thích chi tiết =====
        self.hdg_tn(cau_sau_xu_ly[1] if len(cau_sau_xu_ly) > 1 else None, xml)



    def list_answers_tn(self, content, answer_para, xml):
            """Tạo danh sách đáp án TN, bỏ prefix A./B./C./D. và KHÔNG bọc <div class='content'>."""
        
            multiple_choices = []

            for array_para in content:

                choice_html = self.convert_content_to_html(array_para if isinstance(array_para, list) else [array_para])

                # Bỏ prefix A. B. C. D. nếu có (đầu câu)
                choice_html = re.sub(r"^(<[^>]+>)*\s*[A-Za-z][\.\)]\s*", "", choice_html)

                multiple_choices.append(choice_html.strip())
            # Lấy đáp án đúng
            if isinstance(answer_para, list) and len(answer_para) > 0:

                answer_text = answer_para[0].text.strip()
            else:
                answer_text = answer_para.text.strip()

            number_of_answer = [c for c in answer_text if c.isdigit()]
            listanswers = SubElement(xml, 'listanswers')

            for i, choice in enumerate(multiple_choices):

                answer = SubElement(listanswers, 'answer')

                SubElement(answer, 'index').text = str(i)

                content_elem = SubElement(answer, 'content')

                # Không bọc <div> nữa, chỉ giữ nội dung HTML thuần
                content_elem.text = choice

                is_correct = 'TRUE' if str(i + 1) in number_of_answer else 'FALSE'

                SubElement(answer, 'isanswer').text = is_correct

    # ... (các import cần thiết vẫn giữ nguyên trong class)


    # Hàm tiện ích loại bỏ thẻ HTML
    import re
    def strip_html(self, html_text):
        # Loại bỏ tất cả thẻ <...>
        text = re.sub(r'<[^>]+>', '', html_text)
        # Loại bỏ các khoảng trắng thừa
        text = text.strip()
        return text

    def hdg_tn(self, array_hdg, xml: Element):
        """
        Hướng dẫn giải TN, giữ HTML (ảnh/table)
        - Nếu có hướng dẫn chi tiết thì thêm explainquestion
        - Nếu chỉ có đáp án đúng thì không thêm
        """
       
        if not array_hdg:
            return
        # Xóa thẻ explainquestion cũ nếu có
        existing_explain = xml.find('explainquestion')
        if existing_explain is not None:
            xml.remove(existing_explain)
        hdg_raw = ''
        # Ghép nội dung thô từ array_hdg
        if isinstance(array_hdg, list):
            for part in array_hdg:
                if hasattr(part, "text"):
                    hdg_raw += part.text.strip() + " "
                elif isinstance(part, list):
                    for p in part:
                        if hasattr(p, "text"):
                            hdg_raw += p.text.strip() + " "
        # Chuyển sang HTML (giữ nguyên tag ảnh/table)
        hdg_html = self.convert_content_to_html(array_hdg)
        plain = re.sub(r'<[^>]+>', '', hdg_html).strip()
        explain_text = ""
        # Nếu có nội dung giải thích thực sự
        if len(plain) > 4:
            explain_text = hdg_html.strip()
            # --- 1) Bỏ số hoặc chữ đáp án đầu dòng, kể cả khi nó bị bọc trong thẻ HTML ---
            # Ví dụ: "<strong>1</strong><br>" hoặc "<strong>A</strong>:" hoặc "1. " ...
            explain_text = re.sub(
                r'^\s*(?:<[^>]+>\s*)*(?:\d+|[A-Za-z])(?:\s*</[^>]+>\s*)*(?:\s*(?:<br\s*/?>|:|\.|,))?\s*',
                '',
                explain_text,
                flags=re.IGNORECASE | re.UNICODE
            )
            # --- 2) Bỏ tiền tố "Giải thích:" kể cả khi bị bọc trong thẻ ---
            # Ví dụ: "<strong>Giải thích:</strong><br>" hoặc "Giải thích<br>"
            explain_text = re.sub(
                r'^\s*(?:<[^>]+>\s*)*Giải\s*thích\s*[:：]?(?:\s*</[^>]+>\s*)*(?:\s*(?:<br\s*/?>))?\s*',
                '',
                explain_text,
                flags=re.IGNORECASE | re.UNICODE
            ).strip()
            # Chỉ thêm thẻ nếu còn nội dung sau khi làm sạch
            if explain_text:
                SubElement(xml, 'explainquestion').text = explain_text.strip()

    def dang_ds(self, cau_sau_xu_ly, xml, audio):
        """Xử lý dạng Đúng/Sai, tách đúng phần phát biểu và HDG"""
        SubElement(xml, 'typeAnswer').text = '1'
        SubElement(xml, 'typeViewContent').text = '0'
        SubElement(xml, 'template').text = '0'
        import re

        paragraphs = cau_sau_xu_ly[0]

        statements = []

        intro_paras = []
        # ✅ Phân loại phần mở đầu và các phát biểu
        for para in paragraphs:
            if isinstance(para, Paragraph) and re.match(r'^[a-z]\s*[\.\)]', para.text.strip(), re.IGNORECASE):

                statements.append(para)
            else:

                intro_paras.append(para)
        # ✅ Phần mở đầu (mô tả tình huống)
        content_html = self.convert_content_to_html(intro_paras)

        if audio and len(audio[0]) > 8:

            link = audio[0].replace('Audio:', '').strip()

            content_html += f'<audio controls=""><source src="{link}" type="audio/mpeg"></audio>'
        SubElement(xml, 'contentquestion').text = content_html
        # ✅ Danh sách phát biểu a/b/c/d
        listanswers = SubElement(xml, 'listanswers')
        for i, para in enumerate(statements):
            ans_html = self.convert_content_to_html([para])
            # --- Bỏ prefix a) / b. / c) / d) (kể cả có tag HTML) ---
            ans_html = re.sub(
                r'^\s*(<[^>]+>)*\s*([A-Za-z])\s*[\.\)]\s*',
                '',
                ans_html
            )
            # cũng bỏ trường hợp prefix nằm trong thẻ <strong> hoặc <b>
            # ans_html = re.sub(
            #     r'^(<strong>|<b>)?\s*([A-Da-d])[\.\)]\s*(</strong>|</b>)?',
            #     '',
            #     ans_html
            # )
            ans_html = re.sub(
                r'^\s*(?:<[^>]*>)*\s*[A-Za-z]\s*(?:<[^>]*>)*\s*[\.\)]\s*(?:<[^>]*>)*\s*',
                '',
                ans_html,
                flags=re.IGNORECASE
            )
            answer = SubElement(listanswers, 'answer')

            SubElement(answer, 'index').text = str(i)

            SubElement(answer, 'content').text = ans_html

            SubElement(answer, 'isanswer').text = 'FALSE'  # tạm thời FALSE, sẽ cập nhật sau
        # ✅ Lấy chuỗi đáp án đúng/sai (ví dụ: 0111, 1010, ...)
        if len(cau_sau_xu_ly[1]) > 0:
            if isinstance(cau_sau_xu_ly[1][0], list):

                ans_text = cau_sau_xu_ly[1][0][0].text.strip()

            else:

                ans_text = cau_sau_xu_ly[1][0].text.strip()

            for i, ch in enumerate(ans_text):

                if i < len(listanswers):

                    listanswers[i].find('isanswer').text = 'TRUE' if ch == '1' else 'FALSE'
        # ✅ Hướng dẫn giải (HDG)
        if len(cau_sau_xu_ly[1]) > 1:
            flat_hdg = []
            for item in cau_sau_xu_ly[1][1:]:
                if isinstance(item, list):
                    flat_hdg.extend(item)
                else:
                    flat_hdg.append(item)
            hdg_html = self.convert_content_to_html(flat_hdg)
        else:
            hdg_html = ''
        SubElement(xml, 'explainquestion').text = hdg_html

    def dang_dt(self, cau_sau_xu_ly, xml, subject):
        from xml.etree.ElementTree import SubElement
        import re
        from bs4 import BeautifulSoup

        # ===== META =====
        SubElement(xml, 'typeAnswer').text = '5'
        SubElement(xml, 'typeViewContent').text = '0'
        SubElement(xml, 'template').text = '23'

        # ===== HINT =====
        if len(cau_sau_xu_ly) > 1 and isinstance(cau_sau_xu_ly[1], list) and len(cau_sau_xu_ly[1]) > 1:
            hint_html = self.convert_b4_add_dt(cau_sau_xu_ly[1][1])
            SubElement(xml, 'hintQuestion').text = hint_html

        # ===== RAW HTML =====
        raw_html = self.convert_b4_add_dt(cau_sau_xu_ly[0])  # <-- PHẢI KHÔNG CÓ <p>!

        # GAS không có <p>, chỉ có <br> giữa các đoạn
        lines = [ln for ln in raw_html.split('<br>') if True]  # giữ cả dòng rỗng để xử lý logic GAS

        if not lines:
            lines = ['']

        # ===== TITLE =====
        current_title_txt = lines[0]

        # ✅ Giữ nguyên HTML của dòng title (GAS không strip HTML khi xét title)
        # Chỉ kiểm tra plain text để quyết định có dùng title gốc không
        title_plain = BeautifulSoup(current_title_txt, 'html.parser').get_text().strip()

        final_title = ''
        if len(title_plain) > 1:
            final_title = current_title_txt
        else:
            # Trích xuất toàn bộ đáp án để xác định title mặc định
            found_answers = re.findall(r'\[\[(.*?)\]\]', raw_html)
            all_ans = ''.join(found_answers)
            if subject in getattr(self, 'subjects_with_default_titles', set()):
                if any(c.isalpha() for c in all_ans):
                    final_title = 'Điền đáp án thích hợp vào ô trống'
                else:
                    final_title = 'Điền đáp án thích hợp vào ô trống (chỉ sử dụng chữ số, dấu "," và dấu "-")'

        # ===== XỬ LÝ NỘI DUNG VÀ ĐÁP ÁN theo logic GAS =====
        content_lines = []
        answer_lines = []
        check_one_content = False  # flag như GAS

        # Duyệt từ dòng thứ 1 trở đi (sau title)
        for line in lines[1:]:
            is_input = bool(re.search(r'\[\[.*?\]\]', line))
            is_not_empty = len(line.strip()) > 1

            if not is_input and is_not_empty and not check_one_content:
                content_lines.append(line)
            else:
                check_one_content = True
                answer_lines.append(line)

        # Ghép lại
        content_html = '<br>'.join(content_lines)
        answer_html_raw = '<br>'.join(answer_lines)

        # ===== XỬ LÝ ĐÁP ÁN =====
        input_index = 0
        dap_an_dt = []

        def repl(match):
            nonlocal input_index
            answer_text = match.group(1).strip()
            dap_an_dt.append(answer_text)
            input_index += 1
            return (f'<span class="ans-span-second"></span>'
                    f'<input class="can-resize-second" type="text" id="mathplay-answer-{input_index}">')

        answer_html_processed = re.sub(r'\[\[(.*?)\]\]', repl, answer_html_raw)

        # ===== BUILD XML =====
        cq = SubElement(xml, 'contentquestion')

        # --- title ---
        if final_title:
            title_div = SubElement(cq, 'div')
            title_div.set('class', 'title')
            title_div.text = final_title

        # --- content ---
        content_div = SubElement(cq, 'div')
        content_div.set('class', 'content')
        content_div.text = content_html

        # --- answer-input ---
        if answer_html_processed.strip():
            ans_block = SubElement(cq, 'div')
            ans_block.set('class', 'answer-input')
            # GAS: mỗi dòng trong answer_html_processed → một <div class="line">
            for line in answer_html_processed.split('<br>'):
                if line.strip():
                    line_block = SubElement(ans_block, 'div')
                    line_block.set('class', 'line')
                    line_block.text = line

        # ===== LIST ANSWERS =====
        listanswers = SubElement(xml, 'listanswers')
        for i, ans in enumerate(dap_an_dt):
            ans_clean = ans.replace('‘', "'").replace('’', "'").replace('|', '[-]')
            ans_tag = SubElement(listanswers, 'answer')
            SubElement(ans_tag, 'index').text = str(i)
            SubElement(ans_tag, 'content').text = ans_clean
            SubElement(ans_tag, 'isanswer').text = 'TRUE'

        # ===== EXPLAIN =====
        hdg_html = ''
        if len(cau_sau_xu_ly) > 1 and isinstance(cau_sau_xu_ly[1], list) and cau_sau_xu_ly[1]:
            hdg_html = self.convert_b4_add_dt(cau_sau_xu_ly[1][0])
            
            hdg_plain = BeautifulSoup(hdg_html, 'html.parser').get_text().strip()
        else:
            hdg_plain = ''

        exp = SubElement(xml, 'explainquestion')
        if len(hdg_plain) > 4:
            exp.text = hdg_html
        else:
            exp.text = f"Đáp án đúng theo thứ tự là: {', '.join(dap_an_dt)}"


    def dang_tl(self, cau_sau_xu_ly, xml, audio):
            """Xử lý dạng Tự luận, giữ table/ảnh trong content và HDG"""
            SubElement(xml, 'typeAnswer').text = '3'

            SubElement(xml, 'typeViewContent').text = '0'

            SubElement(xml, 'template').text = '0'
            # Content
            content_html = self.convert_content_to_html(cau_sau_xu_ly[0])

            if audio and len(audio[0]) > 8:

                link = audio[0].replace('Audio:', '').strip()

                content_html += f'<audio controls=""><source src="{link}" type="audio/mpeg"></audio>'

            SubElement(xml, 'contentquestion').text = content_html

            # List answers placeholder
            listanswers = SubElement(xml, 'listanswers')

            answer = SubElement(listanswers, 'answer')

            SubElement(answer, 'index').text = '0'

            SubElement(answer, 'content').text = 'REPLACELATER'

            SubElement(answer, 'isanswer').text = 'TRUE'
            # HDG
            hdg_html = self.convert_content_to_html(cau_sau_xu_ly[1]) if len(cau_sau_xu_ly) > 1 else ''

            SubElement(xml, 'explainquestion').text = hdg_html

    def convert_b4_add_dt(self, paragraphs):
        """Trả về HTML giống GAS: không có <p>, chỉ nối bằng <br>"""
        new_children_all = []
        for index, paragraph in enumerate(paragraphs):
            new_children = []
            if isinstance(paragraph, Table):
                html_table = self.convert_table_to_html(paragraph)
                new_children.append(html_table)
            else:
                self.convert_normal_paras(paragraph, index, new_children)
            new_content = "".join(new_children)
            new_children_all.append(new_content)

        # GAS: chỉ thêm <br> nếu có nhiều đoạn
        if len(new_children_all) > 1:
            string_content = '<br>'.join(new_children_all)
            
        else:
            string_content = new_children_all[0] if new_children_all else ''

        # Xử lý math-latex
        import re
        math_latex = re.compile(r"\$[^$]*\$")
        string_content = math_latex.sub(lambda m: f' <span class="math-tex">{m.group()}</span>', string_content)

        return string_content        

    def convert_b4_add(self, paragraphs):
        """Xử lý danh sách paragraph thành HTML (giống GAS ConvertB4Add)"""
        string_content = '<p>'

        for index, paragraph in enumerate(paragraphs):

            new_children = []
            # if paragraph._element.tag.endswith('tbl'):
            #     html_table = self.convert_table_to_html(paragraph)
            #     new_children.append(html_table)
            if isinstance(paragraph, Table):

                html_table = self.convert_table_to_html(paragraph)

                new_children.append(html_table)
            else:
                self.convert_normal_paras(paragraph, index, new_children)
            new_content = "".join(new_children)

            string_content += f"{new_content}<br>"
        # string_content += "</div>"
        string_content += "</p>"
        # Xử lý math-latex: $...$
        import re
        math_latex = re.compile(r"\$[^$]*\$")

        string_content = math_latex.sub(lambda m: f' <span class="math-tex">{m.group()}</span>', string_content)

        return string_content

    def convert_normal_paras(self, paragraph: Paragraph, index, new_children: list):
        """Chuyển 1 paragraph sang HTML, bỏ phần đầu (Câu, HL, A/B/C/D) và giữ format,
        xử lý cả trường hợp các phần đó bị chia nhỏ qua nhiều run."""
        import re
        # ✅ Gom từng run để dò pattern, kể cả khi chia nhỏ
        progressive_text = ""

        content_start_pos = 0

        detected = False

        patterns = []

        if index == 0:

            patterns.append(r"^C[âa]u\s*\d+[\.:]\s*")  # Câu 1:

        patterns.append(r"^HL:\s*")
        patterns.append(r"^([A-Z])\.\s*")
        # Dò dần theo run
        for run in paragraph.runs:
            if detected:
                break
            full_text = run.text or ""
            progressive_text += full_text
            for pat in patterns:
                m = re.match(pat, progressive_text, re.IGNORECASE)
                if m:
                    content_start_pos = m.end()
                    detected = True
                    break
        # ✅ Sau khi có content_start_pos, xử lý như cũ
        html_content = ""
        prev_style = None
        buffer = ""
        current_text_pos = 0
        for run in paragraph.runs:

            full_text = run.text or ""

            text_start = current_text_pos

            text_end = current_text_pos + len(full_text)

            if text_end <= content_start_pos:

                current_text_pos = text_end


                continue
            if text_start < content_start_pos:

                slice_start = content_start_pos - text_start

                segment_text = full_text[slice_start:]

            else:

                segment_text = full_text
            style = (
                bool(run.bold),
                bool(run.italic),
                bool(run.underline),
                bool(getattr(run.font, 'superscript', False)),
                bool(getattr(run.font, 'subscript', False)),
                bool(getattr(run.font, 'strike', False))
            )
            if prev_style is not None and style != prev_style:
                html_content += self.wrap_style(self.escape_html(buffer), prev_style)
                buffer = ""
            buffer += segment_text
            prev_style = style
            current_text_pos = text_end
        if buffer:
            html_content += self.wrap_style(self.escape_html(buffer), prev_style)
        # ✅ Giữ logic thêm ảnh cũ
        # for run in paragraph.runs:
        #     blips = run._element.xpath(
        #         './/*[local-name()="blip" and namespace-uri()="http://schemas.openxmlformats.org/drawingml/2006/main"]'
        #     )
        #     if blips:
        #         try:
        #             rId_nodes = run._element.xpath(
        #                 './/*[local-name()="blip"]/@*[local-name()="embed"]'
        #             )
        #             if rId_nodes:
        #                 rId = rId_nodes[0]
        #                 img_tag = self._make_img_tag_from_rid(rId)
        #                 if img_tag:
        #                     html_content += img_tag
        #         except Exception:
        #             pass

        for run in paragraph.runs:
            # Tìm tất cả drawing elements trong run
            drawings = run._element.findall(
                './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'
            )
            
            for drawing in drawings:
                try:
                    # 1. Lấy rId từ blip
                    blip = drawing.find(
                        './/{http://schemas.openxmlformats.org/drawingml/2006/main}blip'
                    )
                    
                    if blip is not None:
                        # Lấy r:embed attribute
                        rId = blip.get(
                            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed'
                        )
                        
                        if rId:
                            # 2. Lấy kích thước từ Word XML (EMU units)
                            width_emu, height_emu = self.lay_kich_thuoc_tu_word_xml(drawing)
                            
                            # 3. Tạo HTML img tag với kích thước chính xác
                            img_tag = self._make_img_tag_from_rid(rId, width_emu, height_emu)
                            
                            if img_tag:
                                html_content += img_tag
                                
                except Exception as e:
                    print(f"[ERROR] Xử lý ảnh trong run: {e}")
                    import traceback
                    traceback.print_exc()
        # html_content = self.normalize_line_breaks(html_content)   
        html_content = html_content.replace('####', '')     
        new_children.append(html_content.strip())

    def escape_html(self, text):
        """Escape HTML entities"""
        return (text
            .replace('&', '&amp;')
            .replace('<', '<')
            .replace('>', '>')
            .replace('"', '&quot;')
            .replace("'", '&#039;'))

    def prettify_xml(self, elem):
        """Tạo XML đẹp với indentation"""
        rough_string = tostring(elem, encoding='utf-8')
        reparsed = minidom.parseString(rough_string)
        return reparsed.toprettyxml(indent="  ", encoding='UTF-8').decode('utf-8')


    def post_process_xml(self, xml_str):
        """
        Sửa lại hàm post_process_xml:
        - Di chuyển second_correction ra khỏi vòng lặp đầu
        - Thay đổi cách xử lý math-tex để lấy nội dung bên trong span
        - Thêm các regex để unescape các thẻ có attribute như <table class='...'>
        - Một số sửa nhỏ khác để tránh phá hỏng XML quá sớm
        """
        import re
        from xml.dom import minidom
        import html

        # đảm bảo header
        xml_str = xml_str.replace('<?xml version="1.0" ?>', '<?xml version="1.0" encoding="UTF-8"?>')

        # các thay thế cố định (dùng re.escape khi cần)
        correction = {
            'REPLACELATER': '',
            '&lt;br&gt;': '<br>',
            '&lt;br/&gt;': '<br/>',
            '&lt;em&gt;': '<em>',
            '&lt;/em&gt;': '</em>',
            '&lt;u&gt;': '<u>',
            '&lt;/u&gt;': '</u>',
            '&lt;strong&gt;': '<strong>',
            '&lt;/strong&gt;': '</strong>',
            '&lt;/font&gt;': '</font>',
            '&lt;font': '<font',
            '&lt;span': '<span',
            '&lt;/span&gt;': '</span>',
            '&lt;input': '<input',
            '"&gt;': '">',
            '&lt;/div&gt;': '</div>',
            '&lt;div': '<div',
            '&#xD;': '',
            '&lt;label': '<label',
            '&lt;select': '<select',
            '&lt;option': '<option',
            'hidden&gt;': 'hidden>',
            '&lt;/option&gt;': '</option>',
            '&lt;/select&gt;': '</select>',
            '&lt;/label&gt;': '</label>',
            '&quot;': '"',
            '&lt;center&gt;': '<center>',
            '&lt;/center&gt;': '</center>',
            '&lt;p&gt;': '<p>',
            '&lt;/p&gt;': '</p>',
            '&lt;img': '<img',
            ' /&gt;': ' />',
            '/&gt;': '/>',
            '&lt;audio': '<audio',
            '&lt;/audio&gt;': '</audio>',
            '&lt;source': '<source',
            '&lt;blockquote&gt;': '<blockquote>',
            '&lt;/blockquote&gt;': '</blockquote>',
            '&lt;table&gt;': '<table>',
            '&lt;/table&gt;': '</table>',
            '&lt;tr&gt;': '<tr>',
            '&lt;/tr&gt;': '</tr>',
            '&lt;td&gt;': '<td>',
            '&lt;/td&gt;': '</td>',
            '&lt;li&gt;': '<li>',
            '&lt;/li&gt;': '</li>',
            '&lt;i&gt;': '<i>',
            '&lt;/i&gt;': '</i>',
            '&lt;sub&gt;': '<sub>',
            '&lt;/sub&gt;': '</sub>',
            '&lt;sup&gt;': '<sup>',
            '&lt;/sup&gt;': '</sup>',
        }

        # first pass of simple replacements
        for key, val in correction.items():
            xml_str = re.sub(re.escape(key), val, xml_str, flags=re.IGNORECASE)

        # second set of corrections (ensure it is NOT nested inside the previous loop)
        second_correction = {
            '&lt;i&gt;': '<i>',
            '&lt;/i&gt;': '</i>',
            '&lt;u&gt;': '<u>',
            '&lt;/u&gt;': '</u>',
            '&lt;strong&gt;': '<strong>',
            '&lt;/strong&gt;': '</strong>',
            '&lt;sub&gt;': '<sub>',
            '&lt;/sub&gt;': '</sub>',
            '&lt;sup&gt;': '<sup>',
            '&lt;/sup&gt;': '</sup>',
        }
        for key, val in second_correction.items():
            xml_str = re.sub(re.escape(key), val, xml_str, flags=re.IGNORECASE)

        # === XỬ LÝ CÁC THẺ CÓ ATTR (ví dụ: &lt;table class='...'&gt;) ===
        tags_with_attrs = [
            'table', 'tr', 'td', 'th', 'tbody', 'thead', 'tfoot',
            'img', 'div', 'span', 'p', 'sup', 'sub', 'input', 'label',
            'select', 'option', 'audio', 'source', 'blockquote', 'li', 'center', 'font'
        ]
        for tag in tags_with_attrs:
            xml_str = re.sub(r'&lt;(' + tag + r'\b)', r'<\1', xml_str, flags=re.IGNORECASE)
            xml_str = re.sub(r'&lt;\/(' + tag + r')\s*&gt;', r'</\1>', xml_str, flags=re.IGNORECASE)

        # chuyển các thực thể HTML phổ biến sang ký tự thật (an toàn hơn là unescape toàn bộ)
        xml_str = html.unescape(xml_str)

        # === XỬ LÝ MATHLATEX ===
        def clean_mathlatex(match):
            inner = match.group(1)
            inner = (
                inner
                .replace('<strong>', '')
                .replace('</strong>', '')
                .replace('<i>', '')
                .replace('</i>', '')
                .replace('<u>', '')
                .replace('</u>', '')
                .replace('<br>', '')
                .replace('<br/>', '')
                .replace('%', '\\%')
                .replace('\\frac', '\\dfrac')
            )
            return inner

        xml_str = re.sub(
            r'<span\s+class=["\']math-tex["\']\s*>(.*?)</span>',
            clean_mathlatex,
            xml_str,
            flags=re.DOTALL | re.IGNORECASE
        )

        # === LÀM ĐẸP LẠI XML ===
        try:
            xml_str = minidom.parseString(xml_str.encode('utf-8')).toprettyxml(indent="  ", encoding="UTF-8").decode("utf-8")
        except Exception:
            pass

        # === LƯU FILE ===
        # file_name = "docXML.xml"
        # if "<itemDocuments>" in xml_str:
        #     file_name = "docHL.xml"
        # try:
        #     with open(file_name, "w", encoding="utf-8") as f:
        #         f.write(xml_str)
        # except Exception:
        #     pass

        return xml_str
