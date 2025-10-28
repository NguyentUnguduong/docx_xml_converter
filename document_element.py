from typing import Any, Dict, List
from docx import Document
from docx.document import Document as DocumentType
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table, _Row, _Cell
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.shape import InlineShape
from io import BytesIO


class DocumentElement:
    """Wrapper class to handle different document elements uniformly"""
    
    ELEMENT_TYPES = {
        'PARAGRAPH': 'PARAGRAPH',
        'TEXT': 'TEXT',
        'TABLE': 'TABLE',
        'INLINE_IMAGE': 'INLINE_IMAGE'
    }
    
    def __init__(self, element, element_type=None):
        self.element = element
        self._element_type = element_type
    
    def get_type(self):
        """Get element type"""
        if self._element_type:
            return self._element_type
        
        if isinstance(self.element, Paragraph):
            return self.ELEMENT_TYPES['PARAGRAPH']
        elif isinstance(self.element, Run):
            return self.ELEMENT_TYPES['TEXT']
        elif isinstance(self.element, Table):
            return self.ELEMENT_TYPES['TABLE']
        elif isinstance(self.element, InlineShape):
            return self.ELEMENT_TYPES['INLINE_IMAGE']
        else:
            return 'UNKNOWN'


def get_num_children(element: Any) -> int:
    """Get number of children in element"""
    if isinstance(element, DocumentElement):
        element = element.element
    
    if isinstance(element, Paragraph):
        # Count runs and inline shapes
        count = len(element.runs)
        # Add inline shapes count
        if hasattr(element, '_element') and hasattr(element._element, 'xpath'):
            inline_shapes = element._element.xpath('.//w:drawing')
            count += len(inline_shapes)
        return count
    elif isinstance(element, Table):
        return len(element.rows)
    elif isinstance(element, _Row):
        return len(element.cells)
    elif isinstance(element, _Cell):
        return len(element.paragraphs)
    
    return 0


def get_child(element: Any, index: int) -> Any:
    """Get child element at index"""
    if isinstance(element, DocumentElement):
        element = element.element
    
    if isinstance(element, Paragraph):
        # Get runs and inline shapes
        runs = element.runs
        
        # Check if we need to include inline shapes
        if hasattr(element, '_element'):
            # Get all children (runs and inline shapes)
            children = []
            for child in element._element:
                if child.tag.endswith('r'):  # Run element
                    # Find corresponding run object
                    for run in runs:
                        if run._element == child:
                            children.append(DocumentElement(run, 'TEXT'))
                            break
                elif child.tag.endswith('drawing'):  # Drawing/Image
                    # Create inline shape wrapper
                    try:
                        blip = child.xpath('.//a:blip', namespaces={
                            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                        })[0]
                        rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        image_part = element.part.related_parts[rId]
                        
                        # Get image dimensions
                        extent = child.xpath('.//wp:extent', namespaces={
                            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
                        })[0]
                        width = int(extent.get('cx')) / 9525  # Convert EMU to pixels (approx)
                        height = int(extent.get('cy')) / 9525
                        
                        img_wrapper = {
                            'type': 'INLINE_IMAGE',
                            'blob': image_part.blob,
                            'width': width,
                            'height': height
                        }
                        children.append(DocumentElement(img_wrapper, 'INLINE_IMAGE'))
                    except:
                        pass
            
            if index < len(children):
                return children[index]
        
        # Fallback to runs only
        if index < len(runs):
            return DocumentElement(runs[index], 'TEXT')
    
    elif isinstance(element, Table):
        if index < len(element.rows):
            return element.rows[index]
    elif isinstance(element, _Row):
        if index < len(element.cells):
            return element.cells[index]
    elif isinstance(element, _Cell):
        if index < len(element.paragraphs):
            return element.paragraphs[index]
    
    return None


def get_element_type(element: Any) -> str:
    """Get element type"""
    if isinstance(element, DocumentElement):
        return element.get_type()
    elif isinstance(element, Paragraph):
        return 'PARAGRAPH'
    elif isinstance(element, Run):
        return 'TEXT'
    elif isinstance(element, Table):
        return 'TABLE'
    elif isinstance(element, dict) and element.get('type') == 'INLINE_IMAGE':
        return 'INLINE_IMAGE'
    
    return 'UNKNOWN'


def get_text(element: Any) -> str:
    """Get text from element"""
    if isinstance(element, DocumentElement):
        element = element.element
    
    if isinstance(element, Paragraph):
        return element.text
    elif isinstance(element, Run):
        return element.text
    elif isinstance(element, _Cell):
        return element.text
    
    return ''


def get_blob(element: Any) -> bytes:
    """Get blob from image element"""
    if isinstance(element, DocumentElement):
        element = element.element
    
    if isinstance(element, dict) and element.get('type') == 'INLINE_IMAGE':
        return element['blob']
    
    return b''


def get_width(element: Any) -> int:
    """Get width of element"""
    if isinstance(element, DocumentElement):
        element = element.element
    
    if isinstance(element, dict) and element.get('type') == 'INLINE_IMAGE':
        return int(element.get('width', 100))
    
    return 100


def get_height(element: Any) -> int:
    """Get height of element"""
    if isinstance(element, DocumentElement):
        element = element.element
    
    if isinstance(element, dict) and element.get('type') == 'INLINE_IMAGE':
        return int(element.get('height', 100))
    
    return 100


def get_bytes(blob: bytes) -> bytes:
    """Get bytes from blob"""
    return blob


def get_text_attribute_indices(text_element: Any) -> List[int]:
    """
    Get text attribute indices where formatting changes
    Returns list of positions where formatting attributes change
    """
    if isinstance(text_element, DocumentElement):
        text_element = text_element.element
    
    if not isinstance(text_element, Run):
        return [0]
    
    # For a Run object, formatting is uniform across the entire text
    # But we return [0] to indicate the start position
    # In Google Apps Script, this returns positions where formatting changes
    # In python-docx, each Run has uniform formatting
    text = text_element.text or ''
    return [0] if text else []


def get_attributes(text_element: Any, position: int) -> Dict:
    """Get formatting attributes at position"""
    if isinstance(text_element, DocumentElement):
        text_element = text_element.element
    
    if not isinstance(text_element, Run):
        return {}
    
    return {
        'BOLD': text_element.bold,
        'ITALIC': text_element.italic,
        'UNDERLINE': text_element.underline is not None
    }


def process_style_tinhoc(text_element: Any) -> str:
    """
    Process text style for Tin học
    Returns HTML formatted text with style tags
    """
    if isinstance(text_element, DocumentElement):
        text_element = text_element.element
    
    if not isinstance(text_element, Run):
        return ''
    
    text = text_element.text or ''
    
    # Apply formatting
    if text_element.bold:
        text = f'<strong>{text}</strong>'
    if text_element.italic:
        text = f'<i>{text}</i>'
    if text_element.underline:
        text = f'<u>{text}</u>'
    
    return text


def get_num_rows(table: Any) -> int:
    """Get number of rows in table"""
    if isinstance(table, Table):
        return len(table.rows)
    return 0


def get_row(table: Any, index: int) -> Any:
    """Get row at index"""
    if isinstance(table, Table) and index < len(table.rows):
        return table.rows[index]
    return None


def get_num_cells(row: Any) -> int:
    """Get number of cells in row"""
    if isinstance(row, _Row):
        return len(row.cells)
    return 0


def get_cell(row: Any, index: int) -> Any:
    """Get cell at index"""
    if isinstance(row, _Row) and index < len(row.cells):
        return row.cells[index]
    return None


# ============================================
# Document Loading Helper
# ============================================

def load_document(file_path: str) -> DocumentType:
    """
    Load a Word document from file path
    
    Args:
        file_path: Path to the .docx file
    
    Returns:
        Document object
    """
    return Document(file_path)


def get_paragraphs_from_document(doc: DocumentType) -> List[Paragraph]:
    """
    Get all paragraphs from document including those in tables
    
    Args:
        doc: Document object
    
    Returns:
        List of paragraphs
    """
    paragraphs = []
    
    for element in doc.element.body:
        if isinstance(element, CT_P):
            # Regular paragraph
            para = Paragraph(element, doc)
            paragraphs.append(para)
        elif isinstance(element, CT_Tbl):
            # Table - extract paragraphs from cells
            table = Table(element, doc)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        paragraphs.append(para)
    
    return paragraphs


# ============================================
# Enhanced Text Processing
# ============================================

def extract_runs_with_images(paragraph: Paragraph) -> List[DocumentElement]:
    """
    Extract all runs and inline images from a paragraph in order
    
    Args:
        paragraph: Paragraph object
    
    Returns:
        List of DocumentElement objects (runs and images)
    """
    elements = []
    
    if not hasattr(paragraph, '_element'):
        return [DocumentElement(run, 'TEXT') for run in paragraph.runs]
    
    for child in paragraph._element:
        if child.tag.endswith('r'):  # Run element
            # Find corresponding run object
            for run in paragraph.runs:
                if run._element == child:
                    elements.append(DocumentElement(run, 'TEXT'))
                    break
        
        elif child.tag.endswith('drawing'):  # Drawing/Image
            try:
                # Extract image information
                blip = child.xpath('.//a:blip', namespaces={
                    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
                })[0]
                rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                image_part = paragraph.part.related_parts[rId]
                
                # Get image dimensions
                extent = child.xpath('.//wp:extent', namespaces={
                    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
                })[0]
                width = int(extent.get('cx')) / 9525  # Convert EMU to pixels
                height = int(extent.get('cy')) / 9525
                
                img_data = {
                    'type': 'INLINE_IMAGE',
                    'blob': image_part.blob,
                    'width': width,
                    'height': height
                }
                elements.append(DocumentElement(img_data, 'INLINE_IMAGE'))
            except Exception as e:
                # Skip if image extraction fails
                print(f"Warning: Failed to extract image: {e}")
                pass
    
    return elements